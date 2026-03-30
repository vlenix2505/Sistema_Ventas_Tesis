"""
Genera el documento explicativo del sistema de recomendación en formato DOCX.
Ejecutar: python generar_explicacion.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def titulo(doc, texto, nivel=1):
    p = doc.add_heading(texto, level=nivel)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        if nivel == 1:
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif nivel == 2:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif nivel == 3:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def parrafo(doc, texto, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p

def parrafo_mixto(doc, partes):
    """partes = lista de (texto, bold, italic, color)"""
    p = doc.add_paragraph()
    for texto, bold, italic, color in partes:
        run = p.add_run(texto)
        set_font(run, bold=bold, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p

def viñeta(doc, texto, nivel=0):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(texto)
    set_font(run, size=11)
    p.paragraph_format.left_indent = Inches(0.25 * (nivel + 1))
    p.paragraph_format.space_after = Pt(3)
    return p

def cuadro_gris(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'E8F0FE')
    p._p.get_or_add_pPr().append(shading)
    run = p.add_run(texto)
    set_font(run, size=10, italic=True, color=(0x1F, 0x49, 0x7D))
    return p

def separador(doc):
    doc.add_paragraph()

def tabla_simple(doc, encabezados, filas, col_widths=None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    # Encabezado
    for i, h in enumerate(encabezados):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10, color=(0xFF, 0xFF, 0xFF))
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '2E74B5')
        cell._tc.get_or_add_tcPr().append(shading)
    # Filas
    for idx, fila in enumerate(filas):
        row = t.add_row()
        fill = 'FFFFFF' if idx % 2 == 0 else 'EAF2FB'
        for j, val in enumerate(fila):
            cell = row.cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), fill)
            cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("\n\n\nSISTEMA DE RECOMENDACIÓN HÍBRIDO")
set_font(r, size=24, bold=True, color=(0x1F, 0x49, 0x7D))

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("Guía Explicativa Completa\nCómo funciona, qué hace y por qué funciona")
set_font(r2, size=14, italic=True, color=(0x40, 0x40, 0x40))

doc.add_paragraph()
p_org = doc.add_paragraph()
p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_org.add_run("Empresa distribuidora de alimentos — Sector HORECA — Perú")
set_font(r3, size=12, color=(0x70, 0x70, 0x70))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "Contenido del Documento", 1)
secciones = [
    ("1.", "¿Qué problema resuelve este sistema?", "3"),
    ("2.", "¿Cómo funciona por dentro? — Vista general", "4"),
    ("3.", "Los datos: de dónde vienen y qué contienen", "5"),
    ("4.", "Filtrado Colaborativo (SVD) — aprender de otros clientes", "7"),
    ("5.", "Filtrado por Contenido (CBF) — similitud entre productos", "10"),
    ("6.", "Las Reglas de Negocio — urgencia, rotación y novedad", "12"),
    ("7.", "Cómo se combinan los 5 componentes", "15"),
    ("8.", "Las Métricas — ¿cómo saber si el sistema funciona bien?", "16"),
    ("9.", "Cómo ejecutar el sistema completo", "19"),
    ("10.", "Glosario de términos técnicos", "21"),
]
for num, texto, pag in secciones:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r_num = p.add_run(f"{num}  ")
    set_font(r_num, bold=True, color=(0x2E, 0x74, 0xB5))
    r_txt = p.add_run(texto)
    set_font(r_txt)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 1 — El Problema
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "1. ¿Qué problema resuelve este sistema?", 1)

parrafo(doc,
    "Imagina que eres el dueño de una empresa que distribuye alimentos a restaurantes, "
    "hoteles y cafeterías en Lima, Arequipa, Cusco y Piura. Tienes 950 productos en "
    "catálogo, 800 clientes y un almacén lleno de mercadería que se mueve a distintas velocidades.")

separador(doc)
titulo(doc, "Los tres dolores de cabeza del negocio", 2)

cuadro_gris(doc,
    "Problema 1: Las mermas (pérdidas por vencimiento)\n"
    "Tienes 200 litros de leche que vencen en 5 días. Si no los vendes hoy, los tiras a "
    "la basura. Necesitas saber AHORA a qué clientes les gustan los lácteos para llamarles "
    "y ofrecerles un descuento. El sistema hace exactamente eso.")

separador(doc)

cuadro_gris(doc,
    "Problema 2: El stock parado (productos que nadie compra)\n"
    "Tienes 500 kg de harina especial que nadie ha pedido en 3 meses. El dinero está "
    "literalmente parado en el almacén. El sistema identifica qué clientes tienen más "
    "probabilidad de comprar ese producto y los pone al frente de la lista.")

separador(doc)

cuadro_gris(doc,
    "Problema 3: Los productos nuevos (nadie sabe que existen)\n"
    "Incorporaste una nueva marca de aceite de oliva premium hace 2 semanas. Solo 3 "
    "clientes la han probado. El sistema identifica qué clientes tienen el perfil adecuado "
    "para adoptarla y los recomienda primero.")

separador(doc)
parrafo(doc,
    "La solución es un sistema inteligente que, para cada cliente, genera una lista "
    "personalizada de productos considerando simultáneamente sus gustos históricos Y "
    "las prioridades del negocio (urgencia, rotación, novedad).",
    bold=False)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 2 — Vista General
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "2. ¿Cómo funciona por dentro? — Vista general", 1)

parrafo(doc,
    "El sistema combina 5 ingredientes para calcular un puntaje (score) para cada "
    "combinación cliente-producto. El producto con mayor puntaje va primero en la lista.")

separador(doc)

tabla_simple(doc,
    ["Ingrediente", "¿Qué hace?", "Peso en la fórmula"],
    [
        ("Filtrado Colaborativo (CF)", "Aprende de lo que compraron otros clientes similares a ti", "40%"),
        ("Filtrado por Contenido (CBF)", "Compara características de productos que ya compraste", "15%"),
        ("Score de Urgencia", "Sube el puntaje a productos que vencen pronto", "20%"),
        ("Score de Rotación", "Sube el puntaje a productos que nadie está comprando", "15%"),
        ("Score de Novedad", "Sube el puntaje a productos recién llegados al catálogo", "10%"),
    ],
    col_widths=[2.0, 3.0, 1.5]
)

parrafo(doc, "La fórmula final es:", bold=True)

cuadro_gris(doc,
    "score_final = (0.40 × qué tanto le gusta al cliente)\n"
    "            + (0.15 × qué tan parecido es a sus compras anteriores)\n"
    "            + (0.20 × qué tan urgente es el vencimiento)\n"
    "            + (0.15 × qué tan parado está en el almacén)\n"
    "            + (0.10 × qué tan nuevo es en el catálogo)")

parrafo(doc,
    "El resultado es un número entre 0 y 1. El sistema devuelve los 10 productos "
    "con mayor score para ese cliente (top-10).")

separador(doc)
titulo(doc, "¿Por qué esta combinación?", 2)
parrafo(doc,
    "Solo con inteligencia artificial pura (CF + CBF), el sistema recomendaría "
    "lo que el cliente siempre compra — sin ayudar con las mermas ni el stock parado. "
    "Solo con reglas de negocio, recomendaría productos urgentes a clientes que jamás "
    "los comprarían. La combinación de ambos enfoques permite personalización Y "
    "gestión de inventario al mismo tiempo.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 3 — Los Datos
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "3. Los datos: de dónde vienen y qué contienen", 1)

parrafo(doc,
    "El sistema trabaja con datos sintéticos (generados artificialmente) que imitan "
    "el comportamiento real de una distribuidora de alimentos. Son 4 tablas que se "
    "conectan entre sí, como una base de datos de un negocio real.")

separador(doc)
titulo(doc, "Tabla 1: Clientes (800 registros)", 2)
tabla_simple(doc,
    ["Campo", "Ejemplo", "Qué significa"],
    [
        ("cliente_id", "CLI_000001", "Código único del cliente"),
        ("rubro_cliente", "Restaurante", "Tipo de negocio"),
        ("subrubro_1", "Criolla", "Especialidad principal"),
        ("subrubro_2", "Fusión", "Especialidad secundaria"),
        ("sede_cliente", "Lima", "Ciudad del cliente"),
    ],
    col_widths=[1.5, 1.5, 3.5]
)

titulo(doc, "Tabla 2: Productos (950 únicos, 1,799 con sedes)", 2)
tabla_simple(doc,
    ["Campo", "Ejemplo", "Qué significa"],
    [
        ("producto_id", "PROD_000234", "Código único del producto"),
        ("categoria_producto", "Lácteos", "Categoría (10 categorías en total)"),
        ("precio_unitario", "12.50", "Precio de venta al cliente (S/)"),
        ("COSTO_UNITARIO", "8.20", "Costo de compra para la empresa (S/)"),
        ("stock", "48", "Unidades disponibles en almacén"),
        ("dias_en_stock", "45", "Días que lleva en el almacén"),
        ("fecha_ingreso_catalogo", "2024-11-01", "Cuándo se agregó al catálogo"),
        ("fecha_min_caducidad", "2025-01-08", "Fecha de vencimiento más próxima del lote"),
    ],
    col_widths=[2.0, 1.5, 3.0]
)

titulo(doc, "Tabla 3: Ventas (7,950 registros)", 2)
tabla_simple(doc,
    ["Campo", "Ejemplo", "Qué significa"],
    [
        ("venta_id", "VTA_000001", "Código único de la orden de compra"),
        ("cliente_id", "CLI_000001", "Quién compró"),
        ("fecha_venta", "2024-10-15", "Cuándo fue la compra"),
        ("monto_total", "350.80", "Total de la orden en soles"),
    ],
    col_widths=[1.5, 1.5, 3.5]
)

titulo(doc, "Tabla 4: Detalle de Venta (36,003 registros)", 2)
tabla_simple(doc,
    ["Campo", "Ejemplo", "Qué significa"],
    [
        ("venta_id", "VTA_000001", "A qué orden pertenece esta línea"),
        ("producto_id", "PROD_000234", "Qué producto se compró"),
        ("cantidad_producto", "5", "Cuántas unidades se compraron"),
        ("subtotal", "62.50", "Total de esa línea (5 × 12.50)"),
    ],
    col_widths=[1.5, 1.5, 3.5]
)

separador(doc)
titulo(doc, "El Dataset ML: juntando todo", 2)
parrafo(doc,
    "El notebook 01 une estas 4 tablas y agrega columnas calculadas. "
    "El resultado es un archivo con 290,064 filas y 28 columnas, donde cada fila "
    "representa 'el cliente X compró el producto Y en la fecha Z'. "
    "Ese archivo es el que usa el modelo para aprender.")

cuadro_gris(doc,
    "Columnas adicionales calculadas (las más importantes):\n\n"
    "• dias_para_vencer: cuántos días le quedan al producto antes de vencer\n"
    "• rotacion_diaria: cuántas unidades por día se venden en promedio\n"
    "• baja_rotacion: 1 si está en el 25% de productos menos vendidos, 0 si no\n"
    "• dias_desde_ingreso: cuántos días lleva el producto en catálogo\n"
    "• margen_pct: ganancia porcentual = (precio - costo) / precio")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 4 — Filtrado Colaborativo (SVD)
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "4. Filtrado Colaborativo (SVD) — aprender de otros clientes", 1)

titulo(doc, "La idea intuitiva", 2)
parrafo(doc,
    "¿Alguna vez te recomendaron una película porque 'a personas con gustos similares '  "
    "a los tuyos también les gustó'? Eso es el Filtrado Colaborativo.")

cuadro_gris(doc,
    "Ejemplo concreto:\n\n"
    "El Restaurante A (criolla, Lima) siempre compra: leche, queso, mantequilla, crema de leche.\n"
    "El Restaurante B (criolla, Lima) compra: leche, queso, mantequilla.\n\n"
    "El Restaurante B nunca ha comprado crema de leche. Pero como compra casi lo mismo "
    "que el Restaurante A, el sistema infiere: 'probablemente al Restaurante B también "
    "le gustaría la crema de leche' → la recomienda.")

separador(doc)
titulo(doc, "La matriz de interacciones", 2)
parrafo(doc,
    "Para aplicar esta idea a 800 clientes y 950 productos, el sistema construye "
    "una tabla gigante (llamada 'matriz de interacciones') donde:")

viñeta(doc, "Cada fila es un cliente")
viñeta(doc, "Cada columna es un producto")
viñeta(doc, "Cada celda contiene un número que indica 'cuánto' compró ese cliente ese producto")
viñeta(doc, "La mayoría de celdas están vacías (0), porque ningún cliente compra los 950 productos")

separador(doc)
cuadro_gris(doc,
    "La matriz tiene 800 × 950 = 760,000 celdas, pero solo ~22,913 tienen valor.\n"
    "Eso es una densidad de apenas el 3%: el 97% de las celdas está vacío.\n"
    "A esto se le llama 'matriz sparse' (dispersa).")

separador(doc)
titulo(doc, "El peso por recencia", 2)
parrafo(doc,
    "No todas las compras valen igual. Una compra de hace 2 años dice menos sobre "
    "los gustos actuales que una de la semana pasada. Por eso, el valor en la celda "
    "no es simplemente 'cuántas unidades compró', sino:")

cuadro_gris(doc,
    "valor = cantidad_comprada × e^(-días_desde_la_compra / 180)\n\n"
    "Donde 180 es la 'vida media': una compra de hace 180 días vale la mitad\n"
    "que una compra de hoy. Una compra de hace 360 días vale solo el 25%.")

parrafo(doc,
    "Ejemplo: El cliente compró 10 unidades hace 90 días → valor = 10 × e^(-90/180) = 10 × 0.61 = 6.1\n"
    "Si hubiera comprado esas 10 unidades ayer → valor = 10 × e^(-1/180) ≈ 9.95",
    italic=True, color=(0x40, 0x40, 0x40))

separador(doc)
titulo(doc, "¿Qué es SVD? (Descomposición en Valores Singulares)", 2)
parrafo(doc,
    "El problema con la matriz de 760,000 celdas es que es enorme y llena de vacíos. "
    "SVD es una técnica matemática que 'comprime' esa matriz en algo mucho más pequeño "
    "y manejable.")

cuadro_gris(doc,
    "Analogía: piensa en una foto de alta resolución (mucha información, archivo pesado).\n"
    "SVD es como aplicar compresión JPEG: reduce el tamaño manteniendo lo esencial.\n"
    "En lugar de guardar los 760,000 valores, SVD captura los 'patrones de compra' \n"
    "más importantes en solo 150 dimensiones (los 150 'componentes').")

separador(doc)
parrafo(doc, "¿Cómo funciona técnicamente (explicado simple)?", bold=True)
parrafo(doc,
    "SVD descompone la gran matriz en dos matrices más pequeñas:")
viñeta(doc, "Una matriz de CLIENTES: 800 filas × 150 columnas — cada fila es el 'perfil' de un cliente en 150 dimensiones")
viñeta(doc, "Una matriz de PRODUCTOS: 150 columnas × 950 filas — cada columna es el 'perfil' de un producto")

parrafo(doc,
    "Para recomendar al cliente X, multiplicamos su perfil (1 × 150) por la matriz de "
    "productos (150 × 950) y obtenemos 950 scores en milisegundos. ¡Es muy rápido!")

separador(doc)
cuadro_gris(doc,
    "Resultado del SVD en este proyecto:\n\n"
    "• 150 componentes capturan el 62.8% de la varianza total\n"
    "• Eso significa: 150 dimensiones explican el 62.8% del comportamiento de compra\n"
    "• Es un resultado excelente considerando que la matriz tiene 97% de ceros")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 5 — Filtrado por Contenido (CBF)
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "5. Filtrado por Contenido (CBF) — similitud entre productos", 1)

titulo(doc, "La idea intuitiva", 2)
parrafo(doc,
    "Si siempre compras queso gouda, mantequilla y crema de leche, "
    "probablemente también te gusten el queso edam o la crema ácida. "
    "¿Por qué? Porque son productos muy similares en características.")

cuadro_gris(doc,
    "El Filtrado por Contenido no mira a otros clientes. Solo mira los productos:\n"
    "'Este cliente compra productos de tipo X → busca otros productos parecidos a X'")

separador(doc)
titulo(doc, "¿Cómo se define la 'similitud' entre productos?", 2)
parrafo(doc,
    "Cada producto se convierte en un vector de números que describe sus características:")

tabla_simple(doc,
    ["Característica", "Tipo", "Ejemplo (Leche entera 1L)"],
    [
        ("Categoría", "One-Hot (0/1)", "Lácteos = 1, Bebidas = 0, Carnes = 0 ..."),
        ("Rango de precio", "One-Hot (0/1)", "Precio bajo = 1, medio = 0, alto = 0"),
        ("Rubro principal del cliente típico", "One-Hot (0/1)", "Restaurante = 1, Hotel = 0 ..."),
        ("Precio unitario (normalizado)", "Número 0-1", "0.23 (en escala 0 a 1)"),
        ("Margen porcentual", "Número 0-1", "0.34"),
    ],
    col_widths=[2.2, 1.5, 2.8]
)

parrafo(doc,
    "El resultado es un vector de ~30 números por producto. La 'distancia' entre "
    "dos vectores mide qué tan parecidos son los productos.")

separador(doc)
titulo(doc, "La similitud coseno", 2)
parrafo(doc,
    "Para medir qué tan parecidos son dos vectores, usamos la 'similitud coseno'. "
    "Es el ángulo entre los dos vectores:")

cuadro_gris(doc,
    "• similitud = 1.0: los productos son idénticos en características\n"
    "• similitud = 0.8: muy similares (como leche entera y leche descremada)\n"
    "• similitud = 0.3: algo parecidos (como leche y queso — misma categoría)\n"
    "• similitud = 0.0: completamente distintos (como leche y desinfectante)")

separador(doc)
titulo(doc, "¿Cómo se usa en la recomendación?", 2)
parrafo(doc,
    "Para cada cliente, el sistema mira sus últimas 10 compras y calcula el "
    "promedio de similitud de esas compras con cada producto del catálogo:")

viñeta(doc, "Si el cliente compró: queso, mantequilla, crema de leche, leche")
viñeta(doc, "Se calcula: similitud(queso, leche evaporada) + similitud(mantequilla, leche evaporada) + ...")
viñeta(doc, "Si el promedio es alto → la leche evaporada es parecida a lo que ya compra → sube en la lista")

parrafo(doc,
    "Nota: El CBF solo tiene peso del 15% porque puede caer en el problema del "
    "'más de lo mismo': siempre recomienda lo que el cliente ya compra, sin sorpresas. "
    "El CF (40%) es más potente porque descubre productos que el cliente no conoce aún.",
    italic=True, color=(0x60, 0x60, 0x60))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 6 — Reglas de Negocio
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "6. Las Reglas de Negocio — urgencia, rotación y novedad", 1)

parrafo(doc,
    "Los componentes CF y CBF son 'ciegos al negocio': solo miran el historial de compras. "
    "Las reglas de negocio inyectan inteligencia operativa: la urgencia del inventario, "
    "el stock parado y los nuevos productos.")

separador(doc)
titulo(doc, "6.1 Score de Urgencia (peso 20%) — productos por vencer", 2)

parrafo(doc, "¿Qué es?", bold=True)
parrafo(doc,
    "Un número entre 0 y 1 que indica qué tan urgente es vender el producto antes de "
    "que venza. Se calcula con una curva en forma de 'S' (sigmoide inversa).")

tabla_simple(doc,
    ["Días para vencer", "Score de urgencia", "Interpretación"],
    [
        ("60+ días", "~0.05", "No urgente — no afecta la recomendación"),
        ("30 días", "~0.35", "Moderadamente urgente"),
        ("15 días", "~0.73", "Urgente — sube bastante en la lista"),
        ("7 días", "~0.92", "Muy urgente — casi al tope"),
        ("2 días", "~0.99", "Crítico — debe venderse hoy"),
        ("0 días (vencido)", "Excluido", "El sistema NO recomienda productos vencidos"),
    ],
    col_widths=[1.5, 1.5, 3.5]
)

cuadro_gris(doc,
    "¿Por qué una curva en S y no una línea recta?\n\n"
    "Con una línea recta, la urgencia crecería lentamente de 0 a 1. Con la curva S:\n"
    "• Productos con 60+ días casi no se ven afectados (urgencia ~0)\n"
    "• El 'salto' fuerte ocurre entre 30 y 5 días — justo donde más importa actuar\n"
    "• Esto imita la intuición del almacenero: 'hay que preocuparse cuando queda poco tiempo'")

separador(doc)
titulo(doc, "6.2 Score de Rotación (peso 15%) — productos parados", 2)

parrafo(doc, "¿Qué es?", bold=True)
parrafo(doc,
    "'Rotación diaria' es cuántas unidades se venden por día en promedio. "
    "Un producto con rotación 0.1 significa que se vende 1 unidad cada 10 días. "
    "Un producto con rotación 5.0 se vende 5 unidades por día.")

parrafo(doc,
    "El score de rotación se calcula al revés: los productos que MÁS venden "
    "reciben el score MÁS BAJO. Los que MENOS venden reciben el score MÁS ALTO:")

tabla_simple(doc,
    ["Rotación diaria", "Score de rotación", "Categoría"],
    [
        ("5.0 unidades/día", "0.02", "Alta rotación — no necesita ayuda"),
        ("2.0 unidades/día", "0.35", "Rotación media"),
        ("0.5 unidades/día", "0.75", "Baja rotación — se recomienda activamente"),
        ("0.1 unidades/día", "0.92", "Muy baja rotación — stock parado"),
        ("0.0 unidades/día", "1.00", "Sin movimiento — máxima prioridad de venta"),
    ],
    col_widths=[2.0, 1.5, 3.0]
)

cuadro_gris(doc,
    "Solo los productos en el cuartil inferior (25% con menor rotación) se\n"
    "marcan como 'baja_rotacion'. El sistema los identifica y los empuja hacia arriba\n"
    "en las recomendaciones de clientes que históricamente los comprarían.")

separador(doc)
titulo(doc, "6.3 Score de Novedad (peso 10%) — productos nuevos en catálogo", 2)

parrafo(doc, "¿Qué es?", bold=True)
parrafo(doc,
    "Un producto nuevo es uno que fue incorporado al catálogo hace poco tiempo "
    "(por defecto: menos de 60 días). El score de novedad empieza alto y decae "
    "exponencialmente con el tiempo:")

tabla_simple(doc,
    ["Días desde incorporación", "Score de novedad", "Interpretación"],
    [
        ("1 día", "0.97", "Recién llegado — máxima novedad"),
        ("15 días", "0.61", "Todavía nuevo"),
        ("30 días (1 mes)", "0.37", "Novedad media"),
        ("60 días (2 meses)", "0.14", "Ya no se considera 'nuevo'"),
        ("90+ días", "~0.05", "Sin efecto de novedad"),
    ],
    col_widths=[2.2, 1.5, 2.8]
)

cuadro_gris(doc,
    "Hay una paradoja: si un producto es nuevo, no tiene historial de compras.\n"
    "Eso hace que su score CF sea bajo (nadie lo ha comprado aún).\n"
    "El score de novedad compensa ese déficit para que los productos nuevos\n"
    "aparezcan en las recomendaciones y puedan 'despegar'.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 7 — Cómo se combinan (con datos REALES del dataset — UMBRAL_NOVEDAD=365)
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "7. Cómo se combinan los 5 componentes — Ejemplo real del dataset", 1)

parrafo(doc,
    "Este capítulo usa datos reales extraídos del sistema en ejecución con UMBRAL_NOVEDAD=365. "
    "No son números inventados: son los scores exactos que devuelve la API al consultar "
    "el endpoint /recomendar/dashboard/CLI_074755.")

separador(doc)

# ── Perfil del cliente ────────────────────────────────────────────────────────
titulo(doc, "7.1 El cliente: CLI_074755 — Hostal, Piura", 2)

tabla_simple(doc,
    ["Dato del cliente", "Valor real"],
    [
        ("ID", "CLI_074755"),
        ("Tipo de negocio", "Hotel — Hostal"),
        ("Ciudad", "Piura"),
        ("Período de compras", "Enero 2025 — Febrero 2026 (1 año)"),
        ("Total de líneas de compra", "94  (el cliente más activo del dataset)"),
        ("Productos únicos comprados", "45 de 950 (el 4.7% del catálogo)"),
        ("Ticket promedio por pedido", "S/ 1,741.23"),
        ("Frecuencia de compra", "1.29 pedidos por semana"),
    ],
    col_widths=[2.5, 4.0]
)

separador(doc)
parrafo(doc, "Sus categorías más compradas (unidades totales):", bold=True)

tabla_simple(doc,
    ["Categoría", "Unidades", "Por qué lo compra un hostal"],
    [
        ("Abarrotes",  "469", "Insumos de cocina para el desayuno y servicio diario"),
        ("Verduras",   "155", "Preparaciones frescas para el comedor"),
        ("Frutas",     "108", "Desayunos buffet — alta demanda"),
        ("Limpieza",   "108", "Habitaciones, baños y áreas comunes"),
        ("Lácteos",    "93",  "Leche, yogurt, mantequilla para el desayuno"),
        ("Snacks",     "75",  "Minibar y tienda del hostal"),
        ("Bebidas",    "42",  "Bar y servicio a la habitación"),
        ("Congelados", "14",  "Compra ocasional"),
        ("Carnes",     "10",  "Compra muy baja — se abastece por otro canal"),
    ],
    col_widths=[1.5, 0.9, 4.1]
)

separador(doc)
parrafo(doc, "Productos que el cliente HA COMPRADO del dashboard:", bold=True)

cuadro_gris(doc,
    "De los 15 productos recomendados en el dashboard, solo 1 ha sido comprado antes:\n\n"
    "  PROD_285117 (Frutas): 48 unidades en 1 pedido\n\n"
    "Los 14 restantes son descubrimientos — el modelo los recomienda porque clientes\n"
    "similares (otros hostales de Piura) los compran, o porque el perfil del hostal\n"
    "es compatible con esos productos.")

doc.add_page_break()

# ── Los 3 productos a analizar ────────────────────────────────────────────────
titulo(doc, "7.2 Los 3 productos elegidos para el análisis detallado", 2)

parrafo(doc,
    "Se eligieron 3 productos del top urgentes porque cada uno representa "
    "un caso distinto: uno con historial, uno sin historial pero nuevo y con baja "
    "rotación, y uno también nuevo con baja rotación pero más caro. "
    "Así se puede ver cómo cada componente toma protagonismo en distintos escenarios.")

separador(doc)

tabla_simple(doc,
    ["Campo", "PROD_285117", "PROD_242217", "PROD_087513"],
    [
        ("Categoría",              "Frutas",         "Carnes",           "Lácteos"),
        ("Precio unitario",        "S/ 4.02",        "S/ 127.09",        "S/ 35.92"),
        ("Stock disponible",       "272 unidades",   "237 unidades",     "165 unidades"),
        ("Días para vencer",       "2 días",         "4 días",           "2 días"),
        ("Rotación diaria",        "1.93 unid/día",  "0.29 unid/día",    "0.56 unid/día"),
        ("¿Baja rotación?",        "NO",             "SÍ",               "SÍ"),
        ("Días en catálogo",       "1,536 días",     "192 días",         "148 días"),
        ("¿Es nuevo? (<=365d)",    "NO",             "SÍ",               "SÍ"),
        ("¿Cliente lo ha comprado?","SÍ — 48 unidades","NUNCA",          "NUNCA"),
        ("Score final (API real)", "0.4768  #1",     "0.3387  #2",       "0.3371  #3"),
    ],
    col_widths=[2.0, 1.7, 1.7, 1.7]
)

parrafo(doc,
    "La brecha entre el #1 (0.4768) y el #2 (0.3387) es enorme: 0.1381 puntos. "
    "El CF de PROD_285117 domina completamente. Veamos por qué.",
    italic=True, color=(0x40, 0x40, 0x40))

doc.add_page_break()

# ── Cálculo detallado ─────────────────────────────────────────────────────────
titulo(doc, "7.3 El cálculo completo — score por score", 2)

cuadro_gris(doc,
    "score_final = (CF × 0.40) + (CBF × 0.15) + (Urgency × 0.20) + (Rotation × 0.15) + (Novelty × 0.10)")

separador(doc)

# Producto 1
titulo(doc, "Producto #1: PROD_285117 — Frutas, S/ 4.02, vence en 2 días", 3)
parrafo(doc, "El cliente LO HA COMPRADO antes: 48 unidades. Buena rotación (1.93/día). Lleva 1,536 dias en catálogo.", italic=True)

tabla_simple(doc,
    ["Componente", "Score API", "× Peso", "= Aporte", "Por qué este valor"],
    [
        ("CF — historial + clientes similares", "0.6639", "× 0.40", "= 0.2656",
         "El cliente compró este producto antes Y otros hostales en Piura lo compran mucho. Score muy alto."),
        ("CBF — similitud de características", "0.3138", "× 0.15", "= 0.0471",
         "Frutas es compatible con el hostal, pero el cliente compra más abarrotes. Similitud media."),
        ("Urgency — días para vencer", "0.4667", "× 0.20", "= 0.0933",
         "Vence en 2 días → urgencia alta."),
        ("Rotation — stock parado", "0.4722", "× 0.15", "= 0.0708",
         "Rota 1.93/día → no es baja rotación. Score de rotación medio."),
        ("Novelty — nuevo en catálogo", "0.0000", "× 0.10", "= 0.0000",
         "Lleva 1,536 días en catálogo. No es nuevo (1,536 > 365)."),
        ("SCORE FINAL", "—", "—", "= 0.4768", "El CF solo aporta el 55.7% del score total."),
    ],
    col_widths=[2.0, 0.8, 0.7, 0.7, 2.3]
)
separador(doc)

# Producto 2
titulo(doc, "Producto #2: PROD_242217 — Carnes, S/ 127.09, vence en 4 días", 3)
parrafo(doc, "NUNCA comprado. Baja rotación (0.29/día). 192 dias en catálogo (es NUEVO segun umbral 365). Es urgente, nuevo y baja rotacion a la vez.", italic=True)

tabla_simple(doc,
    ["Componente", "Score API", "× Peso", "= Aporte", "Por qué este valor"],
    [
        ("CF — historial + clientes similares", "0.1203", "× 0.40", "= 0.0481",
         "Nunca comprado. Afinidad inferida: otros hostales en Piura compran carnes ocasionalmente."),
        ("CBF — similitud de características", "0.4212", "× 0.15", "= 0.0632",
         "Carnes es compatible con un hostal que tiene comedor. Score CBF alto."),
        ("Urgency — días para vencer", "0.4337", "× 0.20", "= 0.0867",
         "Vence en 4 días → urgencia alta."),
        ("Rotation — stock parado", "0.9367", "× 0.15", "= 0.1405",
         "0.29 unid/día → muy baja rotacion. Segundo score de rotacion mas alto de los tres."),
        ("Novelty — nuevo en catálogo", "0.0017", "× 0.10", "= 0.0002",
         "192 dias en catalogo → es 'nuevo' (192 <= 365) pero su score de novedad ya decayó mucho."),
        ("SCORE FINAL", "—", "—", "= 0.3387", "Llega al #2 por la combinación CBF + rotation."),
    ],
    col_widths=[2.0, 0.8, 0.7, 0.7, 2.3]
)
separador(doc)

# Producto 3
titulo(doc, "Producto #3: PROD_087513 — Lácteos, S/ 35.92, vence en 2 días", 3)
parrafo(doc, "NUNCA comprado. Baja rotacion (0.56/día). 148 dias en catálogo (es NUEVO). Es urgente, nuevo y baja rotacion a la vez.", italic=True)

tabla_simple(doc,
    ["Componente", "Score API", "× Peso", "= Aporte", "Por qué este valor"],
    [
        ("CF — historial + clientes similares", "0.1211", "× 0.40", "= 0.0484",
         "Nunca comprado. CF similar al de PROD_242217: ambos son desconocidos para el cliente."),
        ("CBF — similitud de características", "0.4367", "× 0.15", "= 0.0655",
         "Lácteos encaja muy bien con un hostal que sirve desayunos. Score CBF alto."),
        ("Urgency — días para vencer", "0.4667", "× 0.20", "= 0.0933",
         "Vence en 2 días → urgencia alta (igual que PROD_285117)."),
        ("Rotation — stock parado", "0.8607", "× 0.15", "= 0.1291",
         "0.56 unid/día → baja rotación. Aunque rota más que PROD_242217, sigue siendo bajo."),
        ("Novelty — nuevo en catálogo", "0.0072", "× 0.10", "= 0.0007",
         "148 dias en catalogo → es 'nuevo' (148 <= 365). Score algo mayor que PROD_242217."),
        ("SCORE FINAL", "—", "—", "= 0.3371", "Queda #3, apenas 0.0016 puntos por debajo del #2."),
    ],
    col_widths=[2.0, 0.8, 0.7, 0.7, 2.3]
)

doc.add_page_break()

# ── Comparación directa ───────────────────────────────────────────────────────
titulo(doc, "7.4 Comparación directa de los 3 productos", 2)

tabla_simple(doc,
    ["Componente", "Peso", "PROD_285117\n(Frutas, comprado)", "PROD_242217\n(Carnes, nuevo)", "PROD_087513\n(Lácteos, nuevo)"],
    [
        ("CF — afinidad",         "40%", "0.2656 ★★★", "0.0481",      "0.0484"),
        ("CBF — similitud",       "15%", "0.0471",      "0.0632",      "0.0655 ★"),
        ("Urgency — vencimiento", "20%", "0.0933 ★",    "0.0867",      "0.0933 ★"),
        ("Rotation — parado",     "15%", "0.0708",      "0.1405 ★",    "0.1291"),
        ("Novelty — nuevo",       "10%", "0.0000",      "0.0002",      "0.0007 ★"),
        ("SCORE FINAL",          "100%", "0.4768 🥇",   "0.3387 🥈",   "0.3371 🥉"),
    ],
    col_widths=[1.8, 0.6, 1.7, 1.7, 1.7]
)
parrafo(doc, "★ = valor mas alto en esa fila", italic=True, color=(0x60, 0x60, 0x60))

separador(doc)

cuadro_gris(doc,
    "POR QUE GANA PROD_285117 A PESAR DE VENCER EN 2 DIAS (igual que PROD_087513):\n\n"
    "La diferencia la hace el CF: 0.2656 vs 0.0484 — una brecha de 0.2172 puntos.\n"
    "Eso equivale a 0.2172 × 0.40 = 0.087 puntos de ventaja solo por CF.\n\n"
    "PROD_087513 intenta compensar con mejor Rotation (0.1291 vs 0.0708 = +0.058)\n"
    "y mejor CBF (0.0655 vs 0.0471 = +0.018), pero la suma de esas mejoras (0.076)\n"
    "no alcanza para superar la ventaja del CF (0.087).\n\n"
    "Conclusion: cuando un producto tiene historial de compra fuerte, es muy dificil\n"
    "que un producto desconocido lo desplace del top, incluso con urgencia maxima.")

separador(doc)

cuadro_gris(doc,
    "POR QUE PROD_242217 (Carnes, S/127.09) SUPERA A PROD_087513 (Lacteos, S/35.92)\n"
    "si ambos son nuevos y nunca comprados:\n\n"
    "El rotation score de PROD_242217 es 0.1405 vs 0.1291 de PROD_087513.\n"
    "Diferencia: 0.0114 puntos. Eso inclina la balanza porque:\n"
    "  PROD_242217 rota 0.29 unid/dia vs 0.56 unid/dia de PROD_087513\n"
    "  → PROD_242217 esta mas parado → el sistema lo empuja con más fuerza.\n\n"
    "El precio alto de PROD_242217 (S/127.09) no penaliza directamente al score;\n"
    "el sistema recomienda por afinidad y urgencia de negocio, no por precio.")

doc.add_page_break()

# ── Triple clasificación ──────────────────────────────────────────────────────
titulo(doc, "7.5 El caso especial: productos que califican en 3 secciones a la vez", 2)

parrafo(doc,
    "PROD_242217 y PROD_087513 tienen algo inusual: califican simultáneamente "
    "como urgentes, baja rotación Y nuevos. ¿En qué sección aparecen?")

tabla_simple(doc,
    ["Producto", "¿Urgente?", "¿Baja rotación?", "¿Nuevo (<=365d)?", "¿Dónde aparece?"],
    [
        ("PROD_285117", "SI (2 días)", "NO",             "NO (1,536 días)", "URGENTES #1"),
        ("PROD_242217", "SI (4 días)", "SI (0.29/día)",  "SI (192 días)",   "URGENTES #2"),
        ("PROD_087513", "SI (2 días)", "SI (0.56/día)",  "SI (148 días)",   "URGENTES #3"),
    ],
    col_widths=[1.5, 1.0, 1.3, 1.5, 1.7]
)

cuadro_gris(doc,
    "El sistema aplica una cascada de prioridad para evitar duplicados:\n\n"
    "  1. Urgentes  (prioridad maxima)\n"
    "  2. Baja rotacion  (excluye los ya en urgentes)\n"
    "  3. Nuevos  (excluye los ya en las dos secciones anteriores)\n\n"
    "PROD_242217 podría aparecer en las 3 secciones, pero solo aparece en URGENTES.\n"
    "Esto es correcto: el argumento de venta más poderoso para el vendedor es\n"
    "'vence en 4 días' — no 'está parado' ni 'es nuevo'.")

separador(doc)

titulo(doc, "7.6 La sección Nuevos — cómo funciona con UMBRAL_NOVEDAD = 365", 2)

parrafo(doc,
    "Con UMBRAL_NOVEDAD = 365, un producto se considera 'nuevo' si fue incorporado "
    "al catálogo hace 365 días o menos. Esto permite que la sección Nuevos tenga "
    "contenido real, ya que el dataset sintético no tiene productos de menos de 60 días.")

tabla_simple(doc,
    ["Producto en sección Nuevos", "Categoría", "Días en catálogo", "¿Por qué aquí y no en urgentes?"],
    [
        ("PROD_815418", "Panadería",  "351 días", "Ya estaba como urgente/baja rotación en otra llamada interna — la cascada lo dejó para Nuevos"),
        ("PROD_669573", "Panadería",  "167 días", "Tambien urgente, pero el pool de urgentes ya tenia 5 productos con scores mas altos"),
        ("PROD_008829", "Verduras",   "266 días", "Idem — score suficiente para Nuevos tras la cascada"),
        ("PROD_174901", "Lácteos",    "251 días", "Idem"),
        ("PROD_602292", "Frutas",     "256 días", "Idem"),
    ],
    col_widths=[1.5, 1.0, 1.2, 3.0]
)

cuadro_gris(doc,
    "Nota sobre score_novelty vs es_nuevo_catalogo:\n\n"
    "El score_novelty (componente de la formula) usa decaimiento exponencial con\n"
    "vida media de 30 dias. Para un producto de 256 dias:\n"
    "  score_novelty = e^(-256/30) = 0.000197  (casi cero)\n\n"
    "Pero es_nuevo_catalogo = True porque 256 <= 365 (el umbral de clasificacion).\n\n"
    "Son dos cosas distintas:\n"
    "  • es_nuevo_catalogo → decide en que SECCION aparece el producto\n"
    "  • score_novelty → cuanto SUBE el score final (muy poco despues de 60 dias)\n\n"
    "En la seccion Nuevos, lo que impulsa el score es principalmente CF y Rotation,\n"
    "no el score de novedad. La novedad es solo la etiqueta de la seccion.")

doc.add_page_break()

# ── Dashboard completo ────────────────────────────────────────────────────────
titulo(doc, "7.7 El dashboard completo — 15 productos reales", 2)

parrafo(doc,
    "Resultado real del endpoint GET /recomendar/dashboard/CLI_074755?top_k=5 "
    "con UMBRAL_NOVEDAD = 365 (valor actual del sistema):")

parrafo(doc, "URGENTES — top 5 (productos que vencen en <= 30 dias):", bold=True, color=(0xC0, 0x00, 0x00))

tabla_simple(doc,
    ["#", "Producto", "Cat.", "Precio", "Dias/vencer", "Rotac./dia", "Nuevo", "Score"],
    [
        ("1", "PROD_285117", "Frutas",   "S/4.02",   "2",  "1.93", "No",  "0.4768"),
        ("2", "PROD_242217", "Carnes",   "S/127.09", "4",  "0.29", "Si",  "0.3387"),
        ("3", "PROD_087513", "Lacteos",  "S/35.92",  "2",  "0.56", "Si",  "0.3371"),
        ("4", "PROD_594009", "Verduras", "S/6.72",   "2",  "0.32", "No",  "0.3355"),
        ("5", "PROD_561973", "Verduras", "S/20.14",  "0",  "0.37", "No",  "0.3353"),
    ],
    col_widths=[0.3, 1.3, 0.9, 0.8, 0.9, 0.9, 0.6, 0.8]
)

parrafo(doc, "BAJA ROTACION — top 5 (excluye los 5 urgentes):", bold=True, color=(0xBF, 0x8F, 0x00))

tabla_simple(doc,
    ["#", "Producto", "Cat.", "Precio", "Dias/vencer", "Rotac./dia", "Nuevo", "Score"],
    [
        ("1", "PROD_962914", "Lacteos",  "S/18.84",  "3",  "0.47", "No",  "0.3291"),
        ("2", "PROD_822162", "Verduras", "S/6.44",   "4",  "0.39", "No",  "0.3248"),
        ("3", "PROD_128264", "Verduras", "S/11.93",  "4",  "0.42", "No",  "0.3232"),
        ("4", "PROD_047797", "Carnes",   "S/77.68",  "3",  "0.36", "No",  "0.3207"),
        ("5", "PROD_285388", "Verduras", "S/14.85",  "0",  "0.46", "No",  "0.3162"),
    ],
    col_widths=[0.3, 1.3, 0.9, 0.8, 0.9, 0.9, 0.6, 0.8]
)

parrafo(doc, "NUEVOS — top 5 (excluye los 10 anteriores, dias_en_catalogo <= 365):", bold=True, color=(0x05, 0x96, 0x69))

tabla_simple(doc,
    ["#", "Producto", "Cat.", "Precio", "Dias/vencer", "Dias catálogo", "Rotac./dia", "Score"],
    [
        ("1", "PROD_815418", "Panaderia", "S/4.96",   "2",  "351", "0.68", "0.3222"),
        ("2", "PROD_669573", "Panaderia", "S/27.11",  "2",  "167", "0.89", "0.3217"),
        ("3", "PROD_008829", "Verduras",  "S/9.98",   "2",  "266", "0.61", "0.3196"),
        ("4", "PROD_174901", "Lacteos",   "S/6.71",   "4",  "251", "0.71", "0.3105"),
        ("5", "PROD_602292", "Frutas",    "S/28.52",  "10", "256", "0.79", "0.3014"),
    ],
    col_widths=[0.3, 1.3, 0.9, 0.8, 0.9, 0.9, 0.7, 0.7]
)

separador(doc)

titulo(doc, "7.8 Los filtros duros — lo que el sistema rechazó antes de calcular", 2)

tabla_simple(doc,
    ["Filtro", "Criterio", "Impacto en CLI_074755 (Piura)"],
    [
        ("Sin stock",       "stock = 0",               "Productos agotados en Piura excluidos automaticamente"),
        ("Vencidos",        "dias_para_vencer < 0",    "Excluidos. dias=0 (vence hoy) SI se incluye — ver PROD_561973 y PROD_285388"),
        ("Sede incorrecta", "sede != Piura",            "Productos de Lima, Arequipa o Cusco excluidos — el cliente es de Piura"),
    ],
    col_widths=[1.4, 1.8, 3.4]
)

cuadro_gris(doc,
    "Atencion: PROD_962914 (Lacteos) aparece en Baja Rotacion con solo 5 unidades en stock.\n"
    "El sistema lo incluye porque stock=5 > 0. El vendedor debe advertir al cliente\n"
    "que solo hay 5 unidades disponibles antes de tomar el pedido.")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 8 — Las Métricas
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "8. Las Métricas — ¿cómo saber si el sistema funciona bien?", 1)

parrafo(doc,
    "Las métricas son como un termómetro: indican la 'temperatura' del modelo. "
    "Aquí las explicamos todas con ejemplos de la vida real.")

separador(doc)
titulo(doc, "8.1 HitRate@10 — ¿el sistema adivina correctamente?", 2)

parrafo(doc, "¿Qué mide?", bold=True)
parrafo(doc,
    "Del total de clientes, ¿en qué porcentaje el próximo producto que compraron "
    "aparece en la lista de 10 recomendaciones del sistema?")

cuadro_gris(doc,
    "Ejemplo concreto:\n\n"
    "Tomamos 100 clientes. Ocultamos su última compra real.\n"
    "Le pedimos al sistema sus 10 mejores recomendaciones para cada uno.\n\n"
    "• Cliente 1: compró leche → el sistema incluyó leche en su top-10 ✓\n"
    "• Cliente 2: compró aceite → el sistema NO incluyó aceite en su top-10 ✗\n"
    "• Cliente 3: compró queso → el sistema incluyó queso en su top-10 ✓\n"
    "... y así para los 100 clientes.\n\n"
    "Si 18 de 100 acertaron: HitRate@10 = 0.18 (18%)")

tabla_simple(doc,
    ["HitRate@10", "Interpretación"],
    [
        ("< 0.05 (5%)", "Sistema malo — no mucho mejor que el azar"),
        ("0.05 – 0.15 (5–15%)", "Aceptable — sistema en proceso de mejora"),
        ("> 0.15 (15%)", "Bueno — target de la tesis"),
        ("> 0.30 (30%)", "Excelente — nivel de Netflix/Amazon"),
    ],
    col_widths=[2.0, 4.5]
)

parrafo(doc,
    "¿Por qué un 15% parece bajo? Porque hay 950 productos. Un sistema al azar "
    "acertaría solo el 1% de las veces (10/950). Alcanzar el 15% significa que el "
    "modelo es 15 veces mejor que elegir al azar.",
    italic=True, color=(0x50, 0x50, 0x50))

separador(doc)
titulo(doc, "8.2 Precision@10 — ¿qué tan relevante es cada recomendación?", 2)

parrafo(doc, "¿Qué mide?", bold=True)
parrafo(doc,
    "De los 10 productos recomendados, ¿cuántos son realmente relevantes para el cliente?")

cuadro_gris(doc,
    "Ejemplo:\n\n"
    "El sistema recomienda 10 productos al cliente.\n"
    "El cliente, en los siguientes 30 días, compra 3 de esos 10.\n\n"
    "Precision@10 = 3/10 = 0.30 (30%)\n\n"
    "Si solo compra 1 de 10: Precision@10 = 0.10 (10%)\n"
    "Si compra 0 de 10: Precision@10 = 0 (el sistema no sirvió de nada)")

parrafo(doc,
    "Target de la tesis: Precision@10 > 0.05 (5%). Parece poco, pero con 950 productos "
    "y compradores que solo necesitan ~5 productos distintos al mes, es un resultado sólido.",
    italic=True, color=(0x50, 0x50, 0x50))

separador(doc)
titulo(doc, "8.3 Coverage — ¿el sistema es diverso o repetitivo?", 2)

parrafo(doc, "¿Qué mide?", bold=True)
parrafo(doc,
    "¿Qué porcentaje del catálogo aparece en alguna recomendación?")

cuadro_gris(doc,
    "Ejemplo:\n\n"
    "El catálogo tiene 950 productos.\n"
    "Después de generar recomendaciones para los 800 clientes:\n"
    "• Si solo 200 productos distintos aparecen en alguna lista → Coverage = 200/950 = 21%\n"
    "• Si 600 productos distintos aparecen → Coverage = 600/950 = 63%\n\n"
    "Target: Coverage > 50%. Significa que al menos la mitad del catálogo\n"
    "está siendo recomendado activamente a alguien.")

parrafo(doc,
    "Un sistema con Coverage bajo tiene un problema llamado 'popularity bias': "
    "siempre recomienda los mismos 50 productos populares e ignora el resto del catálogo. "
    "Eso no ayuda a mover el stock parado.",
    italic=True, color=(0x50, 0x50, 0x50))

separador(doc)
titulo(doc, "8.4 Urgency Coverage — la métrica más importante del negocio", 2)

parrafo(doc, "¿Qué mide?", bold=True)
parrafo(doc,
    "De todos los productos urgentes (que vencen en menos de 30 días), "
    "¿a qué porcentaje el sistema le encontró al menos un cliente compatible para recomendárselo?")

cuadro_gris(doc,
    "Ejemplo:\n\n"
    "Hay 148 productos próximos a vencer en el catálogo.\n\n"
    "• 120 de ellos aparecen en al menos una recomendación a algún cliente\n"
    "• 28 no aparecen en ninguna lista (nadie los comprará antes de que venzan)\n\n"
    "Urgency Coverage = 120/148 = 0.81 (81%)\n\n"
    "Eso significa: el sistema está activamente tratando de mover el 81% de los\n"
    "productos que están en riesgo de convertirse en merma.")

tabla_simple(doc,
    ["Urgency Coverage", "Interpretación para el negocio"],
    [
        ("< 0.30 (30%)", "Sistema ineficaz — la mayoría de merma no se previene"),
        ("0.30 – 0.50 (30–50%)", "Moderado — hay margen de mejora"),
        ("> 0.50 (50%)", "Bueno — target de la tesis"),
        ("> 0.80 (80%)", "Excelente — casi toda la merma potencial se aborda"),
    ],
    col_widths=[2.0, 4.5]
)

separador(doc)
titulo(doc, "8.5 Rotation Coverage — ¿se mueve el stock parado?", 2)

parrafo(doc, "¿Qué mide?", bold=True)
parrafo(doc,
    "De todos los productos con baja rotación, ¿a qué porcentaje el sistema "
    "le encontró un cliente adecuado para recomendárselo?")

cuadro_gris(doc,
    "Ejemplo:\n\n"
    "Hay 549 productos de baja rotación (cuartil inferior).\n\n"
    "• 350 de ellos aparecen en al menos una recomendación\n"
    "• Rotation Coverage = 350/549 = 0.64 (64%)\n\n"
    "Eso significa: el 64% de los productos parados en el almacén están siendo\n"
    "ofrecidos activamente a clientes que probablemente los comprarían.")

separador(doc)
titulo(doc, "Resumen de todas las métricas", 2)
tabla_simple(doc,
    ["Métrica", "Pregunta que responde", "Target", "Interpretación simple"],
    [
        ("HitRate@10", "¿Adivina bien?", "> 15%", "¿Está el producto 'correcto' en la lista?"),
        ("Precision@10", "¿Son relevantes?", "> 5%", "¿Cuántos de los 10 realmente se comprarían?"),
        ("Coverage", "¿Es diverso?", "> 50%", "¿Se recomienda la mitad del catálogo?"),
        ("Urgency Coverage", "¿Reduce merma?", "> 50%", "¿Se atienden los productos urgentes?"),
        ("Rotation Coverage", "¿Mueve stock?", "> 40%", "¿Se atienden los productos parados?"),
    ],
    col_widths=[1.6, 2.0, 0.8, 2.5]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 9 — Cómo ejecutar
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "9. Cómo ejecutar el sistema completo", 1)

titulo(doc, "Requisito previo: Python 3.10 o superior", 2)
parrafo(doc, "Verificar con:")
cuadro_gris(doc, "python --version")

separador(doc)
titulo(doc, "Paso 1 — Instalar dependencias", 2)
cuadro_gris(doc, "pip install -r requirements.txt")
parrafo(doc, "Instala todas las librerías necesarias: pandas, scikit-learn, FastAPI, Streamlit, etc.")

separador(doc)
titulo(doc, "Paso 2 — Generar los datos (si no existen)", 2)
cuadro_gris(doc, "python src/generate_source.py")
parrafo(doc, "Crea los 4 archivos CSV en data/raw/. Tiempo: menos de 30 segundos.")

separador(doc)
titulo(doc, "Paso 3 — Preparar el dataset de ML (si no existe)", 2)
cuadro_gris(doc, "jupyter notebook notebooks/01_dataset.ipynb\n# Ejecutar todas las celdas: Kernel → Restart & Run All")
parrafo(doc, "Une los 4 CSV y crea dataset_ml.csv en data/processed/. Tiempo: 2-5 minutos.")

separador(doc)
titulo(doc, "Paso 4 — Iniciar el Backend (Terminal 1)", 2)
cuadro_gris(doc, "python -m uvicorn api.main:app --reload --port 8000")
parrafo(doc,
    "El servidor arranca y entrena el modelo automáticamente (~1 segundo). "
    "Verifica que funciona abriendo: http://localhost:8000/docs")

separador(doc)
titulo(doc, "Paso 5 — Iniciar el Frontend (Terminal 2)", 2)
cuadro_gris(doc, "python -m streamlit run app/streamlit_app.py")
parrafo(doc, "Abre en el navegador: http://localhost:8501")

separador(doc)
titulo(doc, "Resumen de URLs", 2)
tabla_simple(doc,
    ["Servicio", "URL", "Descripción"],
    [
        ("Backend (API)", "http://localhost:8000", "Servidor de recomendaciones"),
        ("Documentación API", "http://localhost:8000/docs", "Prueba los endpoints visualmente"),
        ("Frontend (Streamlit)", "http://localhost:8501", "Interfaz de vendedor"),
    ],
    col_widths=[2.0, 2.2, 2.3]
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECCIÓN 10 — Glosario
# ══════════════════════════════════════════════════════════════════════════════

titulo(doc, "10. Glosario de términos técnicos", 1)

glosario = [
    ("API (Application Programming Interface)",
     "Es un 'cajero automático' de datos. Le haces una pregunta (request) con un formato "
     "específico y te devuelve la respuesta en un formato definido. En este proyecto, "
     "la API recibe 'dame las recomendaciones para CLI_000001' y devuelve los productos."),
    ("FastAPI",
     "El framework de Python usado para crear la API. Es conocido por ser muy rápido "
     "y por generar automáticamente la documentación interactiva en /docs."),
    ("Streamlit",
     "Una librería de Python que convierte código Python en aplicaciones web visuales "
     "de forma muy simple. El frontend del proyecto está hecho con Streamlit."),
    ("SVD (Singular Value Decomposition)",
     "Técnica matemática de álgebra lineal que 'comprime' una matriz grande en matrices "
     "más pequeñas capturando los patrones más importantes. Similar a la compresión JPEG "
     "pero para datos de comportamiento de usuarios."),
    ("Filtrado Colaborativo (CF)",
     "'Dime con quién andas y te diré qué comprar.' El sistema aprende de los patrones "
     "de compra de todos los clientes para hacer recomendaciones personalizadas."),
    ("Filtrado por Contenido (CBF)",
     "Recomendación basada en las características del producto. Si compraste queso, "
     "probablemente también quieras mantequilla (misma categoría, mismo perfil de precio)."),
    ("Similitud coseno",
     "Medida matemática entre 0 y 1 que indica qué tan parecidos son dos vectores. "
     "Se usa para comparar perfiles de productos. 1 = idénticos, 0 = completamente distintos."),
    ("One-Hot Encoding",
     "Técnica para convertir texto en números. Por ejemplo, la categoría 'Lácteos' se "
     "convierte en un vector [1,0,0,0,0,0,0,0,0,0] donde cada posición representa una categoría."),
    ("MinMaxScaler",
     "Técnica que normaliza los números a un rango de 0 a 1. Por ejemplo, precios de S/1 a S/100 "
     "se convierten en 0.0 a 1.0 para que no dominen el cálculo por ser valores grandes."),
    ("Matriz sparse (dispersa)",
     "Matriz donde la mayoría de los valores son cero. En este proyecto, la matriz de "
     "clientes × productos tiene 97% de ceros porque ningún cliente compra todos los productos."),
    ("Decaimiento exponencial",
     "Función matemática donde el valor decrece rápido al principio y luego lentamente. "
     "Se usa para el score de novedad: los productos muy nuevos tienen score alto, "
     "pero baja rápido en el primer mes y luego se estabiliza cerca de cero."),
    ("Sigmoide inversa",
     "Curva en forma de 'S' al revés: valores altos cuando el input es bajo y viceversa. "
     "Se usa para el score de urgencia: cuando quedan pocos días para vencer (input bajo), "
     "la urgencia sube drásticamente (output alto)."),
    ("Feature engineering",
     "Proceso de crear variables (columnas) nuevas a partir de los datos originales para "
     "ayudar al modelo. Por ejemplo, calcular 'dias_para_vencer' a partir de "
     "'fecha_actual - fecha_caducidad'."),
    ("Top-K",
     "Los K mejores elementos según algún criterio. 'Top-10' significa los 10 productos "
     "con mayor score final. K es configurable: puedes pedir top-5, top-20, etc."),
    ("Pydantic",
     "Librería de Python para validar que los datos tienen el formato correcto. "
     "Si la API espera un número y le mandas texto, Pydantic devuelve un error claro "
     "antes de que cause problemas en el código."),
    ("HitRate@K",
     "Métrica de evaluación: porcentaje de usuarios para quienes el ítem relevante "
     "aparece en el top-K de recomendaciones. Mide la capacidad de 'adivinar' la próxima compra."),
    ("HORECA",
     "Acrónimo de Hoteles, Restaurantes y Cafeterías. Es el sector comercial al que "
     "atiende la distribuidora en este proyecto."),
    ("SKU (Stock Keeping Unit)",
     "Código único que identifica un producto específico. En este proyecto equivale "
     "a producto_id. Una misma empresa puede tener miles de SKUs distintos."),
]

for termino, definicion in glosario:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{termino}: ")
    set_font(r1, bold=True, color=(0x1F, 0x49, 0x7D))
    r2 = p.add_run(definicion)
    set_font(r2)

doc.add_page_break()

# ── Página final ──────────────────────────────────────────────────────────────
p_fin = doc.add_paragraph()
p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fin = p_fin.add_run("\n\n\nFin del documento")
set_font(r_fin, size=13, italic=True, color=(0x70, 0x70, 0x70))

p_sub2 = doc.add_paragraph()
p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub2 = p_sub2.add_run(
    "Sistema de Recomendación Híbrido — Tesis\n"
    "Empresa distribuidora de alimentos | Sector HORECA | Perú"
)
set_font(r_sub2, size=11, color=(0x90, 0x90, 0x90))


# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "docs",
    "explicacion_sistema_recomendacion_v3.docx"
)
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
