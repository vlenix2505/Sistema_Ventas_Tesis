"""
Genera el documento Word con la explicación detallada de recommender.py.
Ejecutar desde la raíz del proyecto:
    python docs/generar_guia_recommender.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = Path(__file__).resolve().parent


# ──────────────────────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def para(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def code(doc, text):
    """Bloque de código en fuente monoespaciada con fondo gris simulado."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            tr.cells[c_idx].text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def margins(doc, top=1, bottom=1, left=1.2, right=1.2):
    s = doc.sections[0]
    s.top_margin    = Inches(top)
    s.bottom_margin = Inches(bottom)
    s.left_margin   = Inches(left)
    s.right_margin  = Inches(right)


def sep(doc):
    doc.add_paragraph("")


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENTO
# ──────────────────────────────────────────────────────────────────────────────

def crear():
    doc = Document()
    margins(doc)

    # Portada
    t = doc.add_heading("recommender.py — Explicación detallada", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph("Cómo está organizado el archivo, qué hace cada parte y en qué orden se ejecuta todo")
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].italic = True
    sep(doc)

    # ── Intro ─────────────────────────────────────────────────────────────────
    heading(doc, "¿Qué es recommender.py?")
    para(doc,
        "recommender.py es el archivo central del sistema. Contiene una sola clase "
        "llamada HybridRecommender, que agrupa toda la lógica del modelo de recomendación: "
        "cómo aprender de los datos, cómo guardar lo aprendido, cómo cargarlo después, "
        "y cómo generar recomendaciones para cada cliente."
    )
    para(doc,
        "Todo lo demás del proyecto (train.py, main.py) solo llama a esta clase. "
        "Si el modelo cambia, solo cambia este archivo."
    )
    sep(doc)

    # ── Estructura general ────────────────────────────────────────────────────
    heading(doc, "Estructura general del archivo")
    para(doc,
        "El archivo está organizado en cuatro bloques bien separados:"
    )
    table(doc,
        headers=["Bloque", "Qué contiene", "Cuándo se ejecuta"],
        rows=[
            ["Parámetros del modelo",
             "Constantes configurables: pesos, umbrales, número de componentes SVD",
             "Al importar el archivo (siempre)"],
            ["__init__()",
             "Declara todas las variables internas vacías. No calcula nada.",
             "Al crear un objeto HybridRecommender()"],
            ["Métodos de entrenamiento\n(fit + 4 sub-métodos)",
             "Aprenden del dataset y guardan los resultados internamente",
             "Solo cuando se ejecuta train.py"],
            ["Métodos de serialización\n(save y load)",
             "Guardan o cargan el modelo entrenado en disco",
             "save: al final de train.py\nload: al arrancar main.py"],
            ["Método de inferencia\n(recommend)",
             "Genera la lista de productos recomendados para un cliente",
             "Por cada request HTTP que llega al servidor"],
        ],
        col_widths=[1.8, 2.8, 2.0],
    )
    sep(doc)

    # ── Parámetros ─────────────────────────────────────────────────────────────
    heading(doc, "Bloque 1 — Parámetros del modelo")
    para(doc,
        "Al inicio del archivo se definen constantes que controlan el comportamiento "
        "del modelo. Son los únicos valores que se cambian si se quiere ajustar el sistema:"
    )
    table(doc,
        headers=["Constante", "Valor", "Qué controla"],
        rows=[
            ["W_CF",               "0.40", "Peso del historial de compras en el score final"],
            ["W_CBF",              "0.15", "Peso de la similitud de contenido"],
            ["W_URGENCY",          "0.20", "Peso de la urgencia por caducidad"],
            ["W_ROTATION",         "0.15", "Peso de la baja rotación"],
            ["W_NOVELTY",          "0.10", "Peso de la novedad en catálogo"],
            ["SVD_COMPONENTS",     "150",  "Cuántas dimensiones latentes usa el SVD"],
            ["TAU_DIAS",           "180",  "Media-vida de recencia (días). Una compra de hace 180 días vale la mitad que hoy"],
            ["TAU_NOVEDAD",        "30",   "Media-vida de novedad. Un producto de 30 días tiene boost = 0.37"],
            ["SIGMA_URGENCY",      "15",   "Pendiente de la curva de urgencia"],
            ["TOPK_CF_CANDIDATES", "500",  "Cuántos candidatos pre-filtra el CF antes de calcular los demás scores"],
            ["TOPK_FINAL",         "10",   "Cuántos productos devuelve la recomendación final"],
            ["N_COMPRAS_RECIENTES","10",   "Cuántas compras recientes del cliente usa el CBF"],
            ["UMBRAL_URGENCIA",    "30",   "Días para considerar un producto 'urgente'"],
            ["UMBRAL_NOVEDAD",     "365",  "Días para considerar un producto 'nuevo en catálogo'"],
        ],
        col_widths=[1.8, 0.6, 4.2],
    )
    sep(doc)

    # ── __init__ ──────────────────────────────────────────────────────────────
    heading(doc, "Bloque 2 — __init__(): preparar el objeto vacío")
    para(doc,
        "Cuando se escribe HybridRecommender(), Python ejecuta __init__(). "
        "Este método no calcula nada. Solo declara las variables internas que el objeto "
        "va a necesitar, todas en None o vacías. Es como preparar los cajones antes de "
        "guardar la ropa."
    )
    para(doc, "Las variables que se inicializan son:", bold=True)
    table(doc,
        headers=["Variable interna", "Qué guardará cuando esté entrenado"],
        rows=[
            ["is_fitted",           "True si el modelo ya fue entrenado o cargado. False al inicio."],
            ["_df",                 "El dataset completo (o una versión reducida con cliente_id y sede)."],
            ["_productos_df",       "Tabla con una fila por producto único y sus atributos."],
            ["_U",                  "Matriz de perfiles de clientes, resultado del SVD. Tamaño: (n_clientes × 150)."],
            ["_Vt",                 "Matriz de perfiles de productos, resultado del SVD. Tamaño: (150 × n_productos)."],
            ["_user_idx",           "Diccionario: cliente_id → número de fila en _U."],
            ["_item_idx",           "Diccionario: producto_id → número de columna en _Vt."],
            ["_idx_item",           "Diccionario inverso: número de columna → producto_id."],
            ["_item_features",      "Matriz de características de contenido de cada producto."],
            ["_cbf_sim",            "Matriz de similitud coseno entre todos los productos. Tamaño: (n_productos × n_productos)."],
            ["_vec_urgency",        "Vector con el score de urgencia pre-calculado para cada producto."],
            ["_vec_novelty",        "Vector con el score de novedad pre-calculado para cada producto."],
            ["_vec_rotation",       "Vector con el score de rotación pre-calculado para cada producto."],
            ["_vec_vencido",        "Vector booleano: 1 si el producto ya venció, 0 si no."],
            ["_stock_by_item",      "Diccionario: producto_id → stock disponible."],
            ["_meta_by_item",       "Diccionario: producto_id → todos sus atributos (precio, días para vencer, etc.)."],
            ["_historial_cliente",  "Diccionario: cliente_id → conjunto de productos que ya compró."],
        ],
        col_widths=[1.8, 4.6],
    )
    sep(doc)

    # ── fit ───────────────────────────────────────────────────────────────────
    heading(doc, "Bloque 3 — fit(): entrenamiento del modelo")
    para(doc,
        "fit() es el método principal de entrenamiento. Se llama una sola vez desde train.py. "
        "Recibe la ruta al archivo dataset_ml.csv, lee los datos y ejecuta cuatro sub-métodos "
        "en orden. Al terminar, todas las variables de __init__() tienen valores."
    )
    code(doc, "rec = HybridRecommender()\nrec.fit('data/processed/dataset_ml.csv')")
    para(doc, "fit() ejecuta estos cuatro pasos internamente:", bold=True)
    sep(doc)

    # ── _build_product_master ────────────────────────────────────────────────
    heading(doc, "Paso 1 dentro de fit() — _build_product_master()", level=2)
    para(doc,
        "El dataset_ml.csv tiene muchas filas por producto (una por cada venta en la que "
        "apareció). Este paso consolida todo en una tabla con una sola fila por producto único."
    )
    para(doc, "¿Cómo elige qué fila conservar?", bold=True)
    para(doc,
        "Si el mismo producto aparece en varias sedes, conserva la fila con mayor stock. "
        "La idea es reflejar la disponibilidad máxima del producto en el sistema."
    )
    para(doc, "¿Qué calcula de nuevo?", bold=True)
    bullet(doc, "días_desde_ingreso: cuántos días pasaron desde que el producto entró al catálogo hasta hoy.")
    bullet(doc, "margen_pct: (precio_unitario - costo_unitario) / precio_unitario. Cuánto gana la empresa por unidad.")
    para(doc,
        "El resultado queda guardado en self._productos_df. Esta tabla es la que usan todos los "
        "pasos siguientes para consultar información de un producto.",
        italic=True,
    )
    sep(doc)

    # ── _build_cf_matrix ─────────────────────────────────────────────────────
    heading(doc, "Paso 2 dentro de fit() — _build_cf_matrix()", level=2)
    para(doc,
        "Este es el paso más importante del entrenamiento. Construye la matriz de interacciones "
        "entre clientes y productos, y la factoriza con SVD para encontrar patrones de compra."
    )

    heading(doc, "Sub-paso A: calcular el peso de recencia de cada compra", level=3)
    para(doc,
        "No todas las compras valen igual. Una compra de hoy dice más sobre los gustos actuales "
        "del cliente que una compra de hace dos años. Por eso se multiplica la cantidad comprada "
        "por un peso que decrece con el tiempo:"
    )
    code(doc, "w_recency = exp( -días_desde_venta / 180 )\ninteraccion = cantidad_producto × w_recency")
    para(doc,
        "Si el cliente compró 10 unidades hace 180 días, su interacción vale 10 × 0.368 = 3.68. "
        "Si las compró hoy, vale 10 × 1.0 = 10.0."
    )

    heading(doc, "Sub-paso B: construir la matriz dispersa R", level=3)
    para(doc,
        "Se agrupan todas las interacciones por (cliente, producto) sumando las compras. "
        "El resultado es una matriz R donde:"
    )
    bullet(doc, "Cada fila es un cliente.")
    bullet(doc, "Cada columna es un producto.")
    bullet(doc, "Cada celda R[cliente, producto] es la suma de todas las interacciones ponderadas por recencia.")
    para(doc,
        "Esta matriz es dispersa (la mayoría de celdas es 0, porque cada cliente "
        "solo compra una fracción de todos los productos del catálogo). Se guarda en "
        "formato comprimido (csr_matrix) para no desperdiciar memoria.",
        italic=True,
    )
    para(doc,
        "En este mismo paso también se guardan tres diccionarios de mapeo que se usan "
        "durante la inferencia:"
    )
    bullet(doc, "_user_idx: convierte un cliente_id (texto) al número de fila en la matriz.")
    bullet(doc, "_item_idx: convierte un producto_id (texto) al número de columna.")
    bullet(doc, "_idx_item: el inverso — de número de columna a producto_id.")
    para(doc,
        "Y también se guarda _historial_cliente: para cada cliente, el conjunto de "
        "productos que alguna vez compró (se usa en el score_cbf)."
    )

    heading(doc, "Sub-paso C: aplicar SVD (Descomposición en Valores Singulares)", level=3)
    para(doc,
        "La matriz R tiene miles de filas y columnas. SVD la comprime en dos matrices más pequeñas "
        "que capturan los patrones más importantes de comportamiento:"
    )
    code(doc, "R  ≈  U × Vᵀ\n\n_U  : (n_clientes × 150)  — perfil latente de cada cliente\n_Vt : (150 × n_productos) — perfil latente de cada producto")
    para(doc,
        "Cada cliente queda representado como un vector de 150 números. Cada producto, igual. "
        "Clientes con gustos similares tendrán vectores similares. "
        "Productos que compran los mismos clientes tendrán vectores similares."
    )
    para(doc,
        "¿Por qué 150? Es el valor de SVD_COMPONENTS. Es un equilibrio entre precisión "
        "(más componentes capturan más detalle) y velocidad (menos componentes, más rápido).",
        italic=True,
    )
    sep(doc)

    # ── _build_cbf_matrix ─────────────────────────────────────────────────────
    heading(doc, "Paso 3 dentro de fit() — _build_cbf_matrix()", level=2)
    para(doc,
        "El Collaborative Filtering (SVD) depende de que el cliente haya comprado algo. "
        "El Content-Based Filtering agrega una segunda señal: la similitud entre productos "
        "según sus atributos. Así, si un producto es similar a lo que el cliente ya compra, "
        "también puede recomendarse aunque no haya un patrón directo en el historial."
    )

    heading(doc, "Sub-paso A: construir el vector de características de cada producto", level=3)
    para(doc, "Cada producto se transforma en un vector numérico con estas features:")
    table(doc,
        headers=["Feature", "Tipo de codificación", "Ejemplo"],
        rows=[
            ["categoria_producto", "One-Hot (una columna por categoría)",
             "Lácteos → [1,0,0,...], Bebidas → [0,1,0,...]"],
            ["precio_bucket",      "One-Hot (bajo / medio / alto)",
             "Precio en el tercio inferior → [1,0,0]"],
            ["rubro_principal",    "One-Hot (rubro más frecuente entre sus compradores)",
             "Restaurante → [0,1,0,...], Bodega → [0,0,1,...]"],
            ["precio_unitario",    "Número normalizado a [0,1]",
             "Precio más bajo del catálogo = 0.0, más alto = 1.0"],
            ["margen_pct",         "Número normalizado a [0,1]",
             "0% margen = 0.0, 100% margen = 1.0"],
        ],
        col_widths=[1.6, 1.9, 3.0],
    )
    para(doc,
        "El resultado es una matriz _item_features de tamaño (n_productos × ~25 columnas). "
        "Cada fila es el vector que describe a un producto.",
        italic=True,
    )

    heading(doc, "Sub-paso B: calcular la matriz de similitud coseno", level=3)
    para(doc,
        "Con los vectores de todos los productos, se calcula qué tan similares son entre sí. "
        "La similitud coseno compara la dirección de dos vectores: si apuntan en la misma "
        "dirección, son similares (valor cercano a 1.0). Si son perpendiculares, no tienen "
        "nada en común (valor 0.0)."
    )
    para(doc,
        "El resultado es _cbf_sim: una matriz cuadrada de (n_productos × n_productos) donde "
        "cada celda [i, j] dice qué tan similares son el producto i y el producto j."
    )
    para(doc,
        "Esta matriz se calcula una sola vez en fit() y se reutiliza en cada llamada a "
        "recommend() sin recalcular nada.",
        italic=True,
    )
    sep(doc)

    # ── _build_business_scores ────────────────────────────────────────────────
    heading(doc, "Paso 4 dentro de fit() — _build_business_scores()", level=2)
    para(doc,
        "Los tres scores de negocio (urgency, rotation, novelty) no dependen del cliente: "
        "son propiedades del producto. Por eso se calculan una sola vez para todos los "
        "productos y se guardan como vectores. En cada llamada a recommend() se consultan "
        "directamente sin recalcular."
    )
    table(doc,
        headers=["Vector", "Fórmula", "Qué mide"],
        rows=[
            ["_vec_urgency",
             "1 / (1 + exp(días_para_vencer / 15))",
             "Qué tan próximo está a vencer. Valores altos = vence pronto."],
            ["_vec_rotation",
             "1 - (rotacion_diaria - min) / (max - min)",
             "Qué tan poco rota el producto. Valores altos = stock parado."],
            ["_vec_novelty",
             "exp(-días_desde_ingreso / 30)",
             "Qué tan nuevo es en el catálogo. Valores altos = recién ingresado."],
            ["_vec_vencido",
             "1 si días_para_vencer < 0, si no 0",
             "Máscara booleana. Los productos ya vencidos se excluyen del resultado."],
        ],
        col_widths=[1.4, 2.3, 2.9],
    )
    sep(doc)

    # ── _build_aux_indices ────────────────────────────────────────────────────
    heading(doc, "Paso 5 dentro de fit() — _build_aux_indices()", level=2)
    para(doc,
        "Un paso pequeño pero importante. Construye dos diccionarios para acceso instantáneo "
        "durante la inferencia:"
    )
    bullet(doc, "_stock_by_item:",
           "producto_id → cantidad en stock. Se usa para el filtro duro 'sin stock'.")
    bullet(doc, "_meta_by_item:",
           "producto_id → todos sus atributos (precio, días para vencer, rotación, etc.). "
           "Se usa para construir la respuesta final.")
    para(doc,
        "Sin estos diccionarios, cada llamada a recommend() tendría que buscar en el DataFrame, "
        "lo cual es mucho más lento.",
        italic=True,
    )
    sep(doc)

    # ── save / load ───────────────────────────────────────────────────────────
    heading(doc, "Bloque 4 — save() y load(): guardar y cargar el modelo")
    para(doc,
        "Una vez que fit() termina, todas las variables internas (_U, _Vt, _cbf_sim, etc.) "
        "están en memoria RAM. Si el servidor se detiene, se pierden. "
        "save() serializa todo eso a un archivo .pkl en disco. load() hace lo inverso."
    )

    heading(doc, "save(path)", level=2)
    para(doc,
        "Empaqueta todas las variables internas en un diccionario y lo guarda en disco "
        "usando joblib (una librería optimizada para arrays de numpy, más rápida que pickle):"
    )
    code(doc,
        "artifacts = {\n"
        "    'U': self._U,\n"
        "    'Vt': self._Vt,\n"
        "    'user_idx': self._user_idx,\n"
        "    'item_idx': self._item_idx,\n"
        "    ... (todos los demás)\n"
        "}\n"
        "joblib.dump(artifacts, path, compress=3)"
    )
    para(doc,
        "compress=3 reduce el tamaño del archivo en disco sin perder información. "
        "Es una compresión sin pérdidas.",
        italic=True,
    )

    heading(doc, "load(path)", level=2)
    para(doc,
        "Crea un objeto HybridRecommender vacío (como __init__) y rellena sus variables "
        "directamente desde el archivo, saltándose todo el entrenamiento:"
    )
    code(doc,
        "rec = HybridRecommender()          # objeto vacío\n"
        "a = joblib.load(path)              # carga el diccionario\n"
        "rec._U    = a['U']                 # restaura las matrices\n"
        "rec._Vt   = a['Vt']\n"
        "rec._cbf_sim = a['cbf_sim']\n"
        "... (todos los demás)\n"
        "rec.is_fitted = True               # marca como listo\n"
        "return rec"
    )
    para(doc,
        "Después de load(), el objeto está listo para llamar a recommend() exactamente "
        "igual que si se hubiera ejecutado fit(). La diferencia: load() tarda menos de "
        "1 segundo; fit() puede tardar varios minutos.",
        italic=True,
    )
    sep(doc)

    # ── recommend ─────────────────────────────────────────────────────────────
    heading(doc, "Bloque 5 — recommend(): generar recomendaciones")
    para(doc,
        "recommend() es el método que se llama por cada request HTTP que llega al servidor. "
        "Recibe un cliente_id y devuelve una lista de productos recomendados. "
        "Ejecuta 7 pasos en secuencia:"
    )
    sep(doc)

    heading(doc, "Paso 1 — Verificar que el cliente existe", level=2)
    para(doc,
        "Antes de calcular cualquier cosa, verifica que el cliente_id esté en _user_idx. "
        "Si no está (cliente sin historial de compras), lanza un error inmediatamente. "
        "Esto evita cálculos innecesarios."
    )
    code(doc, "if cliente_id not in self._user_idx:\n    raise ValueError('Cliente no encontrado')")
    sep(doc)

    heading(doc, "Paso 2 — Calcular score_cf para todos los productos", level=2)
    para(doc,
        "Se busca el vector del cliente en la matriz _U y se multiplica por _Vt. "
        "Esto produce un score_cf para cada uno de los n_productos del catálogo en una sola operación:"
    )
    code(doc,
        "user_vec  = self._U[u_idx]         # vector del cliente: (150,)\n"
        "scores_cf = user_vec @ self._Vt    # multiplicación: (150,) × (150 × n_prod) = (n_prod,)"
    )
    para(doc,
        "El resultado scores_cf es un número por producto. Se normaliza a [0,1] "
        "restando el mínimo y dividiendo entre el rango.",
        italic=True,
    )
    sep(doc)

    heading(doc, "Paso 3 — Pre-filtrar candidatos por score_cf", level=2)
    para(doc,
        "No todos los productos pasan al siguiente paso. Se toman solo los 500 con mayor "
        "score_cf (TOPK_CF_CANDIDATES). Para el resto, no se calculan los demás scores."
    )
    para(doc,
        "¿Por qué 500 y no todos? El catálogo puede tener miles de productos. Calcular "
        "score_cbf para todos sería lento. El pre-filtrado garantiza que los candidatos "
        "más prometedores ya están en el pool.",
        italic=True,
    )
    para(doc,
        "Además, se agregan forzosamente al pool todos los productos urgentes "
        "(score_urgency > 0.3) y los de baja rotación (score_rotation > 0.75), "
        "aunque no estén entre los top-500 por CF. "
        "Esto garantiza que los productos críticos para el negocio siempre tienen "
        "oportunidad de aparecer en el resultado.",
        italic=True,
    )
    sep(doc)

    heading(doc, "Paso 4 — Aplicar filtro por tipo (si se especificó)", level=2)
    para(doc,
        "Los endpoints de la API pueden pedir solo un tipo de productos. "
        "Si se especificó filtro_tipo, el pool se restringe antes de calcular scores:"
    )
    table(doc,
        headers=["filtro_tipo", "Qué hace"],
        rows=[
            ["None (sin filtro)",  "No restringe. Se usan todos los candidatos del paso anterior."],
            ["'urgentes'",         "Solo se conservan productos con días_para_vencer ≤ umbral_urgencia."],
            ["'baja_rotacion'",    "Solo se conservan productos con baja_rotacion == 1."],
            ["'nuevos'",           "Solo se conservan productos con días_desde_ingreso ≤ umbral_novedad."],
        ],
        col_widths=[1.5, 5.0],
    )
    sep(doc)

    heading(doc, "Paso 5 — Calcular score_cbf para los candidatos", level=2)
    para(doc,
        "Para cada producto candidato, se calcula qué tan similar es a las últimas "
        "N_COMPRAS_RECIENTES = 10 compras del cliente. Se usa la matriz _cbf_sim "
        "pre-calculada en fit():"
    )
    code(doc,
        "# Para el candidato i y las últimas 10 compras del cliente:\n"
        "sim_recientes = self._cbf_sim[candidatos_arr][:, compras_recientes]\n"
        "scores_cbf_cand = sim_recientes.mean(axis=1)"
    )
    para(doc,
        "sim_recientes es una submatriz de similitudes. Se promedian las columnas "
        "(una por compra reciente) para obtener un score único por candidato.",
        italic=True,
    )
    sep(doc)

    heading(doc, "Paso 6 — Calcular score_final combinando los 5 scores", level=2)
    para(doc,
        "Con todos los scores listos para los candidatos, se combinan con los pesos configurados:"
    )
    code(doc,
        "score_final = (\n"
        "    0.40 × scores_cf_norm[candidatos]\n"
        "  + 0.15 × scores_cbf_cand\n"
        "  + 0.20 × vec_urgency[candidatos]\n"
        "  + 0.15 × vec_rotation[candidatos]\n"
        "  + 0.10 × vec_novelty[candidatos]\n"
        ")"
    )
    para(doc,
        "Estas son operaciones vectoriales de numpy: se calculan todos los candidatos "
        "al mismo tiempo, no uno por uno. Por eso la inferencia es tan rápida.",
        italic=True,
    )
    sep(doc)

    heading(doc, "Paso 7 — Aplicar filtros duros", level=2)
    para(doc,
        "Algunos productos no deben recomendarse bajo ninguna circunstancia, "
        "independientemente de su score:"
    )
    bullet(doc, "Productos sin stock (stock == 0): no tiene sentido ofrecerlos si no hay disponibilidad.")
    bullet(doc, "Productos vencidos (vec_vencido == 1): no se pueden vender.")
    para(doc,
        "A estos productos se les asigna score_final = -∞ para que queden "
        "automáticamente al final del ordenamiento y nunca aparezcan en el top-K.",
        italic=True,
    )
    sep(doc)

    heading(doc, "Paso 8 — Ordenar y construir el resultado final", level=2)
    para(doc,
        "Se ordenan todos los candidatos por score_final de mayor a menor y se toman "
        "los primeros TOPK_FINAL = 10. Para cada uno se construye un diccionario con "
        "todos los campos que la API va a devolver:"
    )
    table(doc,
        headers=["Campo devuelto", "De dónde sale"],
        rows=[
            ["producto_id",         "_idx_item[j]  — convierte el índice de columna al ID del producto"],
            ["categoria_producto",  "_meta_by_item[producto_id]['categoria_producto']"],
            ["precio_unitario",     "_meta_by_item[producto_id]['precio_unitario']"],
            ["stock",               "_meta_by_item[producto_id]['stock']"],
            ["score_final",         "Valor combinado calculado en el paso 6"],
            ["score_cf",            "Componente CF del candidato"],
            ["score_cbf",           "Componente CBF del candidato"],
            ["score_urgency",       "Del vector _vec_urgency[j]"],
            ["score_rotation",      "Del vector _vec_rotation[j]"],
            ["score_novelty",       "Del vector _vec_novelty[j]"],
            ["dias_para_vencer",    "_meta_by_item[producto_id]['dias_para_vencer']"],
            ["dias_desde_ingreso",  "_meta_by_item[producto_id]['dias_desde_ingreso']"],
            ["es_urgente",          "dias_para_vencer ≤ UMBRAL_URGENCIA (flag para la UI)"],
            ["es_nuevo_catalogo",   "dias_desde_ingreso ≤ UMBRAL_NOVEDAD (flag para la UI)"],
            ["es_baja_rotacion",    "_meta_by_item[producto_id]['baja_rotacion'] == 1 (flag para la UI)"],
            ["rotacion_diaria",     "_meta_by_item[producto_id]['rotacion_diaria']"],
        ],
        col_widths=[1.8, 4.7],
    )
    sep(doc)

    # ── Cómo se conecta con main.py ────────────────────────────────────────────
    heading(doc, "Cómo se conecta recommender.py con main.py")
    para(doc,
        "main.py crea una sola instancia de HybridRecommender que vive en memoria "
        "mientras el servidor está encendido. Todos los requests HTTP comparten esa "
        "misma instancia — no se crea una nueva por cada solicitud."
    )
    code(doc,
        "# En main.py — al arrancar el servidor:\n"
        "recommender = HybridRecommender.load('data/processed/modelo_artifacts.pkl')\n\n"
        "# Por cada request GET /recomendar/{cliente_id}:\n"
        "productos = recommender.recommend(cliente_id, top_k=10)\n"
        "# → devuelve la lista de dicts que construye el paso 8"
    )
    para(doc,
        "Los endpoints de filtrado (proximos-vencer, baja-rotacion, nuevos) son "
        "exactamente iguales, solo cambia el parámetro filtro_tipo que se pasa a recommend():"
    )
    code(doc,
        "# Endpoint /recomendar/proximos-vencer/{cliente_id}:\n"
        "productos = recommender.recommend(cliente_id, filtro_tipo='urgentes')\n\n"
        "# Endpoint /recomendar/baja-rotacion/{cliente_id}:\n"
        "productos = recommender.recommend(cliente_id, filtro_tipo='baja_rotacion')\n\n"
        "# Endpoint /recomendar/nuevos/{cliente_id}:\n"
        "productos = recommender.recommend(cliente_id, filtro_tipo='nuevos')"
    )
    sep(doc)

    # ── Diagrama de flujo ─────────────────────────────────────────────────────
    heading(doc, "Diagrama de flujo completo")
    para(doc,
        "A continuación se resume el ciclo de vida completo de recommender.py "
        "en los dos momentos en que se usa:"
    )
    code(doc,
        "════════════════════════════════════════════════════════\n"
        "  MOMENTO 1: python scripts/train.py  (una vez al día)\n"
        "════════════════════════════════════════════════════════\n"
        "HybridRecommender()\n"
        "  └── __init__()              → variables vacías\n"
        "\n"
        "fit('dataset_ml.csv')\n"
        "  ├── _build_product_master() → _productos_df\n"
        "  ├── _build_cf_matrix()      → _U, _Vt, _user_idx, _item_idx,\n"
        "  │                              _idx_item, _historial_cliente\n"
        "  ├── _build_cbf_matrix()     → _item_features, _cbf_sim\n"
        "  ├── _build_business_scores()→ _vec_urgency, _vec_rotation,\n"
        "  │                              _vec_novelty, _vec_vencido\n"
        "  └── _build_aux_indices()    → _stock_by_item, _meta_by_item\n"
        "\n"
        "save('modelo_artifacts.pkl')  → serializa todo a disco\n"
        "\n"
        "════════════════════════════════════════════════════════\n"
        "  MOMENTO 2: uvicorn api.main:app  (servidor encendido)\n"
        "════════════════════════════════════════════════════════\n"
        "load('modelo_artifacts.pkl')  → restaura todo desde disco\n"
        "  is_fitted = True\n"
        "\n"
        "Por cada request HTTP:\n"
        "recommend(cliente_id)\n"
        "  ├── Paso 1: verificar cliente\n"
        "  ├── Paso 2: scores_cf = _U[cliente] @ _Vt\n"
        "  ├── Paso 3: pre-filtrar top-500 candidatos + urgentes + baja rot\n"
        "  ├── Paso 4: aplicar filtro_tipo si se especificó\n"
        "  ├── Paso 5: score_cbf usando _cbf_sim\n"
        "  ├── Paso 6: score_final = 0.40×CF + 0.15×CBF + 0.20×URG + ...\n"
        "  ├── Paso 7: filtros duros (sin stock, vencidos → score = -∞)\n"
        "  └── Paso 8: ordenar y devolver top-10 con todos los campos"
    )
    sep(doc)

    output = DOCS_DIR / "recommender_explicacion_detallada.docx"
    doc.save(output)
    print(f"Generado: {output}")


if __name__ == "__main__":
    crear()
