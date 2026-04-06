"""
Genera los dos documentos Word de documentación del modelo.
Ejecutar desde la raíz del proyecto:
    python docs/generar_documentacion.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).resolve().parent


# ──────────────────────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_paragraph(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            tr.cells[c_idx].text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def set_page_margins(doc, top=1, bottom=1, left=1.25, right=1.25):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENTO 1: GUÍA SIMPLE
# ──────────────────────────────────────────────────────────────────────────────

def crear_doc_simple():
    doc = Document()
    set_page_margins(doc)

    # Portada
    title = doc.add_heading("Sistema de Recomendación Híbrido", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Guía técnica del modelo: cómo funciona y cómo se conecta con la API")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph("")

    # ── 1. Qué hace el sistema ────────────────────────────────────────────────
    add_heading(doc, "1. ¿Qué hace el sistema?")
    add_paragraph(doc,
        "El sistema es un motor de recomendación de productos para una distribuidora de alimentos. "
        "Su función es responder a la pregunta: dado un cliente específico, ¿qué productos debería "
        "ofrecerle el vendedor en este momento?"
    )
    add_paragraph(doc,
        "La respuesta no es una lista fija. Cambia según:"
    )
    add_bullet(doc, "Lo que el cliente ha comprado antes (historial).")
    add_bullet(doc, "Lo que han comprado clientes similares a él.")
    add_bullet(doc, "El estado actual del inventario: productos por vencer, con poco movimiento, o nuevos.")
    doc.add_paragraph("")

    # ── 2. De dónde viene la información ────────────────────────────────────
    add_heading(doc, "2. De dónde viene la información")
    add_paragraph(doc,
        "El modelo se alimenta de cuatro archivos CSV que representan la base de datos de la distribuidora:"
    )
    add_table(doc,
        headers=["Archivo", "Qué contiene", "Ejemplo de campos"],
        rows=[
            ["clientes.csv",      "Datos de cada cliente",         "cliente_id, rubro, sede"],
            ["productos.csv",     "Catálogo de productos",          "producto_id, categoría, precio, stock, fecha_caducidad"],
            ["ventas.csv",        "Cabecera de cada venta",         "venta_id, cliente_id, fecha_venta"],
            ["detalle_venta.csv", "Qué productos se vendieron",     "venta_id, producto_id, cantidad"],
        ],
        col_widths=[1.5, 2.0, 2.5],
    )
    doc.add_paragraph("")
    add_paragraph(doc,
        "Estos cuatro archivos se combinan en el notebook 01_dataset.ipynb, que los une y genera "
        "el archivo dataset_ml.csv. Ese archivo es la fuente de verdad que consume el modelo."
    )
    doc.add_paragraph("")

    # ── 3. Flujo completo ─────────────────────────────────────────────────────
    add_heading(doc, "3. Flujo completo del sistema")
    add_paragraph(doc,
        "El sistema tiene tres etapas claramente separadas:"
    )

    add_heading(doc, "Etapa 1 — Preparación de datos (una vez o cuando cambien los datos)", level=2)
    add_numbered(doc, "Se ejecuta el notebook 01_dataset.ipynb.")
    add_numbered(doc, "Este notebook lee los 4 CSV, los une y calcula campos derivados "
                      "(días para vencer, rotación diaria, etc.).")
    add_numbered(doc, "El resultado es data/processed/dataset_ml.csv.")

    add_heading(doc, "Etapa 2 — Entrenamiento del modelo (una vez al día)", level=2)
    add_numbered(doc, "Se ejecuta el comando: python scripts/train.py")
    add_numbered(doc, "Este script lee dataset_ml.csv y entrena el modelo completo.")
    add_numbered(doc, "El resultado es data/processed/modelo_artifacts.pkl "
                      "(el modelo guardado en disco).")
    add_paragraph(doc,
        "Este paso tarda varios minutos. Por eso se hace una sola vez al día, "
        "preferiblemente antes del horario de atención.",
        italic=True,
    )

    add_heading(doc, "Etapa 3 — Servicio de recomendaciones (en tiempo real)", level=2)
    add_numbered(doc, "El servidor arranca con: uvicorn api.main:app --port 8000")
    add_numbered(doc, "main.py carga el modelo desde modelo_artifacts.pkl en menos de 1 segundo.")
    add_numbered(doc, "El vendedor selecciona un cliente en la app.")
    add_numbered(doc, "La app llama al endpoint GET /recomendar/{cliente_id}.")
    add_numbered(doc, "El servidor responde con una lista de productos recomendados en menos de 100ms.")
    doc.add_paragraph("")

    # ── 4. Qué hay dentro del modelo ─────────────────────────────────────────
    add_heading(doc, "4. Cómo funciona el modelo por dentro")
    add_paragraph(doc,
        "El modelo combina cinco señales para calcular una puntuación final por producto. "
        "Cada señal captura un aspecto diferente de la situación:"
    )
    add_table(doc,
        headers=["Señal", "Peso", "Qué mide", "Fuente de la información"],
        rows=[
            ["score_cf",       "40%", "Preferencia del cliente según historial de compras propio y de clientes similares",   "Historial de ventas (SVD)"],
            ["score_cbf",      "15%", "Similitud del producto con lo que el cliente ya compró",                               "Atributos del producto"],
            ["score_urgency",  "20%", "Qué tan próximo está a vencer el producto",                                           "Fecha de caducidad"],
            ["score_rotation", "15%", "Qué tan poco se ha vendido el producto (stock parado)",                               "Historial de ventas"],
            ["score_novelty",  "10%", "Qué tan nuevo es el producto en el catálogo",                                         "Fecha de ingreso al catálogo"],
        ],
        col_widths=[1.2, 0.6, 2.8, 1.8],
    )
    doc.add_paragraph("")
    add_paragraph(doc,
        "La puntuación final de cada producto se calcula así:"
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "score_final = 0.40 × score_cf  +  0.15 × score_cbf  +  "
        "0.20 × score_urgency  +  0.15 × score_rotation  +  0.10 × score_novelty"
    )
    run.bold = True
    run.font.name = "Courier New"
    doc.add_paragraph("")
    add_paragraph(doc,
        "Antes de devolver el resultado, se aplican dos filtros que eliminan productos "
        "que no deben recomendarse bajo ninguna circunstancia:"
    )
    add_bullet(doc, "Productos sin stock disponible.")
    add_bullet(doc, "Productos cuya fecha de caducidad ya venció.")
    doc.add_paragraph("")

    # ── 5. Qué hace recommender.py ────────────────────────────────────────────
    add_heading(doc, "5. El archivo recommender.py — el corazón del modelo")
    add_paragraph(doc,
        "recommender.py contiene toda la lógica del modelo en una sola clase: HybridRecommender. "
        "Tiene tres operaciones principales:"
    )
    add_table(doc,
        headers=["Método", "Qué hace", "Cuándo se usa"],
        rows=[
            ["fit(dataset_path)",  "Entrena el modelo leyendo el CSV",             "Solo en scripts/train.py"],
            ["save(path)",         "Guarda el modelo entrenado en disco (.pkl)",    "Al final de scripts/train.py"],
            ["load(path)",         "Carga un modelo ya entrenado desde disco",      "Al arrancar el servidor (main.py)"],
            ["recommend(cliente)", "Genera la lista de productos recomendados",     "Por cada request HTTP"],
        ],
        col_widths=[1.8, 2.5, 2.2],
    )
    doc.add_paragraph("")
    add_paragraph(doc,
        "El método fit() realiza internamente cuatro pasos:"
    )
    add_numbered(doc, "_build_product_master() — construye una tabla con un registro por producto único.")
    add_numbered(doc, "_build_cf_matrix() — crea la matriz de interacciones y aplica SVD para encontrar "
                      "patrones de compra entre clientes y productos.")
    add_numbered(doc, "_build_cbf_matrix() — crea vectores numéricos que describen cada producto "
                      "(categoría, precio, rubro de compradores) y calcula qué tan similares son entre sí.")
    add_numbered(doc, "_build_business_scores() — pre-calcula los tres scores de negocio "
                      "(urgency, rotation, novelty) para todos los productos.")
    doc.add_paragraph("")

    # ── 6. Los endpoints y cómo se conectan ───────────────────────────────────
    add_heading(doc, "6. Los endpoints de la API y su conexión con el modelo")
    add_paragraph(doc,
        "main.py expone cinco endpoints. Todos internamente llaman a recommender.recommend(), "
        "cambiando el parámetro filtro_tipo según el caso de uso:"
    )
    add_table(doc,
        headers=["Endpoint", "filtro_tipo", "Para qué sirve"],
        rows=[
            ["GET /recomendar/{cliente_id}",                None,             "Recomendación híbrida general"],
            ["GET /recomendar/proximos-vencer/{cliente_id}", "urgentes",      "Solo productos por vencer"],
            ["GET /recomendar/baja-rotacion/{cliente_id}",  "baja_rotacion",  "Solo productos con stock parado"],
            ["GET /recomendar/nuevos/{cliente_id}",         "nuevos",         "Solo productos nuevos en catálogo"],
            ["GET /recomendar/dashboard/{cliente_id}",      "los 3 anteriores","Las 3 listas sin productos repetidos"],
        ],
        col_widths=[2.8, 1.4, 2.3],
    )
    doc.add_paragraph("")
    add_paragraph(doc,
        "El endpoint /dashboard llama a recommend() tres veces (una por categoría) y luego aplica "
        "una jerarquía: un producto que es urgente Y tiene baja rotación aparece solo en 'urgentes'. "
        "Esto evita que el vendedor vea el mismo producto en dos secciones."
    )
    doc.add_paragraph("")

    # ── 7. Cómo usar el sistema ────────────────────────────────────────────────
    add_heading(doc, "7. Cómo usar el sistema día a día")

    add_heading(doc, "Primera vez (instalación)", level=2)
    add_numbered(doc, "Instalar dependencias: pip install -r requirements.txt")
    add_numbered(doc, "Ejecutar el notebook 01_dataset.ipynb para generar dataset_ml.csv.")
    add_numbered(doc, "Entrenar el modelo: python scripts/train.py")
    add_numbered(doc, "Arrancar el servidor: uvicorn api.main:app --port 8000")

    add_heading(doc, "Uso diario", level=2)
    add_numbered(doc, "Por la mañana (antes de abrir el sistema): python scripts/train.py")
    add_paragraph(doc,
        "Esto actualiza el modelo con las ventas del día anterior. "
        "Solo es necesario si el dataset_ml.csv se actualizó con nuevas ventas.",
        italic=True,
    )
    add_numbered(doc, "Arrancar el servidor normalmente.")
    add_paragraph(doc,
        "El servidor detecta automáticamente si existe un modelo guardado y lo carga. "
        "Si no existe, entrena y guarda uno nuevo como fallback.",
        italic=True,
    )
    doc.add_paragraph("")

    # ── 8. Estructura de carpetas ─────────────────────────────────────────────
    add_heading(doc, "8. Estructura de carpetas del proyecto")
    add_paragraph(doc,
        "Cada carpeta tiene un rol específico y bien delimitado:"
    )
    add_table(doc,
        headers=["Carpeta / Archivo", "Rol", "¿Se modifica manualmente?"],
        rows=[
            ["data/raw/*.csv",                    "Datos originales de la distribuidora",              "Sí, cuando llegan nuevos datos"],
            ["data/processed/dataset_ml.csv",     "Dataset preparado para el modelo",                  "No — lo genera el notebook 01"],
            ["data/processed/modelo_artifacts.pkl","Modelo entrenado y guardado",                      "No — lo genera scripts/train.py"],
            ["notebooks/01_dataset.ipynb",         "Preparación y limpieza de datos",                  "Solo si cambia la fuente de datos"],
            ["scripts/train.py",                   "Script de entrenamiento manual (una vez al día)",  "No"],
            ["api/recommender.py",                 "Toda la lógica del modelo (fit, save, load, recommend)", "Solo si se ajusta el modelo"],
            ["api/main.py",                        "API FastAPI con los endpoints",                    "Solo si se agregan endpoints"],
            ["api/schemas.py",                     "Estructura de los datos que devuelve la API",      "Solo si cambian los campos"],
            ["app/streamlit_app.py",               "Interfaz visual del sistema",                      "Solo si se modifica el frontend"],
        ],
        col_widths=[2.2, 2.5, 1.7],
    )
    doc.add_paragraph("")
    add_paragraph(doc,
        "Archivos que ya no forman parte del flujo del modelo (pueden eliminarse si existen):",
        italic=True,
    )
    add_bullet(doc, "notebooks/02_modelo_recomendacion.ipynb — su lógica fue centralizada en recommender.py.")
    add_bullet(doc, "notebooks/build_nb.py — utilidad de construcción de notebooks, no parte del sistema.")
    add_bullet(doc, "src/generate_source.py — no relacionado con el modelo.")
    add_bullet(doc, "generar_explicacion.py — reemplazado por esta documentación.")

    output = DOCS_DIR / "modelo_guia_simple.docx"
    doc.save(output)
    print(f"Generado: {output}")


# ──────────────────────────────────────────────────────────────────────────────
# DOCUMENTO 2: GUÍA MATEMÁTICA
# ──────────────────────────────────────────────────────────────────────────────

def crear_doc_matematico():
    doc = Document()
    set_page_margins(doc)

    title = doc.add_heading("Sistema de Recomendación Híbrido", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Fundamentos matemáticos del modelo con ejemplos numéricos")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph("")

    # ── 1. Visión general ─────────────────────────────────────────────────────
    add_heading(doc, "1. Visión general del modelo")
    add_paragraph(doc,
        "El sistema combina cinco componentes en una puntuación única por producto:"
    )
    p = doc.add_paragraph()
    p.add_run(
        "score_final(u, i) = W_CF × score_cf(u,i)  +  W_CBF × score_cbf(u,i)  +\n"
        "                    W_URGENCY × score_urgency(i)  +  W_ROTATION × score_rotation(i)  +\n"
        "                    W_NOVELTY × score_novelty(i)"
    ).font.name = "Courier New"
    add_paragraph(doc, "donde u = cliente, i = producto, y los pesos suman 1.0:")
    add_table(doc,
        headers=["Variable", "Valor", "Componente"],
        rows=[
            ["W_CF",       "0.40", "Collaborative Filtering (SVD implícito)"],
            ["W_CBF",      "0.15", "Content-Based Filtering (similitud coseno)"],
            ["W_URGENCY",  "0.20", "Urgencia por caducidad (sigmoide inversa)"],
            ["W_ROTATION", "0.15", "Baja rotación (inverso normalizado)"],
            ["W_NOVELTY",  "0.10", "Novedad en catálogo (decaimiento exponencial)"],
        ],
        col_widths=[1.2, 0.7, 3.5],
    )
    doc.add_paragraph("")

    # ── 2. Matriz de interacciones ────────────────────────────────────────────
    add_heading(doc, "2. Construcción de la matriz de interacciones implícitas")
    add_paragraph(doc,
        "El punto de partida es construir una matriz R de tamaño (clientes × productos). "
        "Cada celda R[u, i] representa cuánto ha comprado el cliente u del producto i, "
        "ponderado por la recencia de la compra."
    )

    add_heading(doc, "Peso de recencia", level=2)
    add_paragraph(doc, "Las compras recientes valen más que las antiguas. Se aplica un decaimiento exponencial:")
    p = doc.add_paragraph()
    p.add_run("w_recency = exp( -días_desde_venta / TAU_DIAS )").font.name = "Courier New"
    add_paragraph(doc,
        "Con TAU_DIAS = 180 días (media-vida), una compra de hace 6 meses vale la mitad "
        "que una compra de hoy."
    )
    add_paragraph(doc, "Ejemplo numérico:", bold=True)
    add_table(doc,
        headers=["Compra", "Días transcurridos", "w_recency", "Cantidad", "Interacción = cantidad × w_recency"],
        rows=[
            ["Hoy",           "0",   "exp(-0/180)   = 1.000", "5 unidades", "5.00"],
            ["Hace 180 días", "180", "exp(-180/180) = 0.368", "5 unidades", "1.84"],
            ["Hace 360 días", "360", "exp(-360/180) = 0.135", "5 unidades", "0.68"],
        ],
        col_widths=[1.1, 1.4, 1.8, 1.2, 2.0],
    )
    doc.add_paragraph("")
    add_paragraph(doc,
        "Si el cliente compró el mismo producto varias veces, las interacciones se suman "
        "para obtener un único valor R[u, i] que resume toda la relación cliente-producto."
    )
    doc.add_paragraph("")

    # ── 3. SVD — Collaborative Filtering ─────────────────────────────────────
    add_heading(doc, "3. Collaborative Filtering — Descomposición SVD")
    add_paragraph(doc,
        "La Descomposición en Valores Singulares (SVD, por sus siglas en inglés) factoriza "
        "la matriz R en tres matrices más pequeñas:"
    )
    p = doc.add_paragraph()
    p.add_run("R  ≈  U × S × Vᵀ").font.name = "Courier New"
    add_table(doc,
        headers=["Matriz", "Dimensiones", "Qué representa"],
        rows=[
            ["U",  "(n_clientes × k)", "Perfil latente de cada cliente. Cada fila es el 'gusto' del cliente en k dimensiones."],
            ["S",  "(k × k)",          "Importancia de cada dimensión latente (valores singulares)."],
            ["Vᵀ", "(k × n_productos)", "Perfil latente de cada producto. Cada columna describe el producto en k dimensiones."],
        ],
        col_widths=[0.8, 1.5, 3.8],
    )
    add_paragraph(doc,
        "El número k = 150 (SVD_COMPONENTS). En la práctica se usa TruncatedSVD de scikit-learn, "
        "que calcula solo los k componentes más importantes sin necesitar la matriz completa.",
        italic=True,
    )
    doc.add_paragraph("")

    add_heading(doc, "Cálculo del score_cf", level=2)
    add_paragraph(doc,
        "Para predecir qué tan compatible es el cliente u con el producto i:"
    )
    p = doc.add_paragraph()
    p.add_run("score_cf(u, i)  =  U[u, :]  ·  Vᵀ[:, i]   (producto punto)").font.name = "Courier New"
    add_paragraph(doc, "Ejemplo simplificado con k = 3 dimensiones:", bold=True)
    add_paragraph(doc,
        "Supón que el cliente A tiene perfil U[A] = [0.8, 0.2, 0.5] "
        "(le gustan mucho las categorías 1 y 3, poco la 2)."
    )
    add_paragraph(doc,
        "Y que el producto X tiene perfil Vᵀ[X] = [0.9, 0.1, 0.6]."
    )
    p = doc.add_paragraph()
    p.add_run(
        "score_cf(A, X) = 0.8×0.9 + 0.2×0.1 + 0.5×0.6 = 0.72 + 0.02 + 0.30 = 1.04"
    ).font.name = "Courier New"
    add_paragraph(doc,
        "Luego este valor se normaliza al rango [0, 1] dividiéndolo entre el máximo y "
        "restando el mínimo de todos los productos. El producto con mayor producto punto "
        "obtiene score_cf = 1.0."
    )
    doc.add_paragraph("")

    # ── 4. CBF — Content-Based Filtering ─────────────────────────────────────
    add_heading(doc, "4. Content-Based Filtering — Similitud coseno")
    add_paragraph(doc,
        "Cada producto se describe como un vector numérico con sus características. "
        "Este vector tiene tres tipos de features:"
    )
    add_bullet(doc, "Categoría del producto → codificación One-Hot (ej. Lácteos = [1,0,0,...], Bebidas = [0,1,0,...]).")
    add_bullet(doc, "Rango de precio → One-Hot (bajo / medio / alto según percentiles p33 y p66).")
    add_bullet(doc, "Rubro principal de sus compradores → One-Hot (ej. Restaurante, Bodega, etc.).")
    add_bullet(doc, "Precio unitario y margen → normalizados a [0,1] con MinMaxScaler.")
    doc.add_paragraph("")

    add_heading(doc, "Similitud coseno entre dos productos", level=2)
    add_paragraph(doc,
        "La similitud coseno mide el ángulo entre dos vectores. "
        "Si el ángulo es 0° (misma dirección), la similitud es 1.0 (productos idénticos). "
        "Si el ángulo es 90°, la similitud es 0.0 (sin nada en común)."
    )
    p = doc.add_paragraph()
    p.add_run(
        "sim_coseno(A, B) = (A · B) / (||A|| × ||B||)"
    ).font.name = "Courier New"
    add_paragraph(doc, "Ejemplo numérico:", bold=True)
    add_paragraph(doc,
        "Producto X: vector = [1, 0, 0, 0.8, 0.3]  (Lácteo, precio alto, margen 0.3)\n"
        "Producto Y: vector = [1, 0, 0, 0.7, 0.4]  (Lácteo, precio alto, margen 0.4)\n"
        "Producto Z: vector = [0, 1, 0, 0.2, 0.1]  (Bebida, precio bajo, margen 0.1)"
    )
    p = doc.add_paragraph()
    p.add_run(
        "sim(X, Y) ≈ 0.99  →  muy similares (misma categoría y rango)\n"
        "sim(X, Z) ≈ 0.11  →  poco similares (categorías distintas)"
    ).font.name = "Courier New"
    doc.add_paragraph("")

    add_heading(doc, "Cálculo del score_cbf para un cliente", level=2)
    add_paragraph(doc,
        "El score_cbf del candidato i para el cliente u se calcula como el promedio de "
        "similitud entre i y las últimas N_COMPRAS_RECIENTES = 10 compras del cliente:"
    )
    p = doc.add_paragraph()
    p.add_run(
        "score_cbf(u, i) = promedio( sim_coseno(i, c)  para c en historial_reciente(u) )"
    ).font.name = "Courier New"
    add_paragraph(doc,
        "Si el cliente compró principalmente Lácteos y el candidato es un Lácteo, "
        "score_cbf será alto. Si el candidato es Bebidas, score_cbf será bajo."
    )
    doc.add_paragraph("")

    # ── 5. Score de urgencia ──────────────────────────────────────────────────
    add_heading(doc, "5. Score de urgencia — Sigmoide inversa")
    add_paragraph(doc,
        "La urgencia mide qué tan próximo está un producto a vencer. "
        "Se usa una función sigmoide invertida: cuando días_para_vencer → 0, "
        "el score → 1.0 (máxima urgencia). Cuando días_para_vencer es grande, el score → 0.0."
    )
    p = doc.add_paragraph()
    p.add_run(
        "score_urgency(i) = 1 / ( 1 + exp( días_para_vencer(i) / SIGMA_URGENCY ) )"
    ).font.name = "Courier New"
    add_paragraph(doc, "Con SIGMA_URGENCY = 15 días:", italic=True)
    add_table(doc,
        headers=["Días para vencer", "score_urgency", "Interpretación"],
        rows=[
            ["0 días",   "1 / (1 + exp(0/15))  = 0.500", "Vence hoy — urgencia media"],
            ["5 días",   "1 / (1 + exp(5/15))  = 0.422", "Vence en 5 días"],
            ["15 días",  "1 / (1 + exp(15/15)) = 0.269", "Vence en 15 días"],
            ["30 días",  "1 / (1 + exp(30/15)) = 0.119", "Vence en 30 días — urgencia baja"],
            ["90 días",  "1 / (1 + exp(90/15)) = 0.002", "Vence en 3 meses — sin urgencia"],
            ["< 0 días", "0.000 (excluido)",             "Ya venció — se elimina de la recomendación"],
        ],
        col_widths=[1.4, 2.5, 2.5],
    )
    doc.add_paragraph("")

    # ── 6. Score de rotación ──────────────────────────────────────────────────
    add_heading(doc, "6. Score de rotación — Inverso normalizado")
    add_paragraph(doc,
        "La rotación_diaria de un producto es cuántas unidades se venden por día desde "
        "su primera aparición en el historial. Un producto con baja rotación necesita "
        "ser impulsado, por lo que recibe un score alto (score_rotation → 1.0)."
    )
    p = doc.add_paragraph()
    p.add_run(
        "score_rotation(i) = 1  -  ( rotacion_diaria(i) - rot_min ) / ( rot_max - rot_min )"
    ).font.name = "Courier New"
    add_paragraph(doc, "Ejemplo numérico:", bold=True)
    add_paragraph(doc,
        "Supón que en el catálogo la rotación mínima es 0.1 u/día y la máxima es 5.0 u/día:"
    )
    add_table(doc,
        headers=["Producto", "Rotación diaria", "score_rotation", "Interpretación"],
        rows=[
            ["Producto A", "0.1 u/día (mínima)", "1 - (0.1-0.1)/(5.0-0.1) = 1.00", "Stock totalmente parado"],
            ["Producto B", "2.5 u/día (media)",  "1 - (2.5-0.1)/(5.0-0.1) = 0.51", "Rotación media"],
            ["Producto C", "5.0 u/día (máxima)", "1 - (5.0-0.1)/(5.0-0.1) = 0.00", "Producto estrella"],
        ],
        col_widths=[1.2, 1.6, 2.5, 2.2],
    )
    doc.add_paragraph("")

    # ── 7. Score de novedad ────────────────────────────────────────────────────
    add_heading(doc, "7. Score de novedad — Decaimiento exponencial")
    add_paragraph(doc,
        "Los productos nuevos en el catálogo tienen el problema del 'cold-start': "
        "como nadie los ha comprado, el score_cf no los favorece. "
        "El score_novelty les da un impulso que decrece con el tiempo."
    )
    p = doc.add_paragraph()
    p.add_run(
        "score_novelty(i) = exp( -días_desde_ingreso(i) / TAU_NOVEDAD )"
    ).font.name = "Courier New"
    add_paragraph(doc, "Con TAU_NOVEDAD = 30 días (media-vida):", italic=True)
    add_table(doc,
        headers=["Días desde ingreso", "score_novelty", "Interpretación"],
        rows=[
            ["0 días",   "exp(-0/30)   = 1.000", "Recién ingresado — máximo boost"],
            ["30 días",  "exp(-30/30)  = 0.368", "1 mes — boost moderado"],
            ["60 días",  "exp(-60/30)  = 0.135", "2 meses — boost pequeño"],
            ["90 días",  "exp(-90/30)  = 0.050", "3 meses — casi sin efecto"],
            ["365 días", "exp(-365/30) = 0.000", "1 año — sin boost de novedad"],
        ],
        col_widths=[1.6, 1.8, 3.0],
    )
    doc.add_paragraph("")

    # ── 8. Score final y pre-filtrado ─────────────────────────────────────────
    add_heading(doc, "8. Combinación final y proceso de inferencia")
    add_paragraph(doc,
        "Cuando la API recibe una solicitud para el cliente u, el método recommend() "
        "ejecuta los siguientes pasos:"
    )
    add_numbered(doc,
        "Paso 1 — score_cf para todos los productos: se multiplica el vector de perfil "
        "del cliente U[u] (tamaño 150) por la matriz Vᵀ (150 × n_productos). "
        "Resultado: un score_cf por producto."
    )
    add_numbered(doc,
        "Paso 2 — Pre-filtrado: se toman los 500 productos con mayor score_cf como candidatos "
        "(TOPK_CF_CANDIDATES = 500). Esto reduce el cálculo posterior."
    )
    add_numbered(doc,
        "Paso 3 — Garantía de cobertura: se agregan al pool todos los productos urgentes "
        "(score_urgency > 0.3) y los de baja rotación (score_rotation > 0.75), "
        "aunque no estén entre los top-500 por CF."
    )
    add_numbered(doc,
        "Paso 4 — score_cbf: para cada candidato, se calcula el promedio de similitud "
        "coseno con las últimas 10 compras del cliente."
    )
    add_numbered(doc,
        "Paso 5 — score_final: se combinan las 5 señales con los pesos definidos."
    )
    add_numbered(doc,
        "Paso 6 — Filtros duros: se eliminan del resultado productos sin stock o vencidos "
        "(su score_final se fija en -∞)."
    )
    add_numbered(doc,
        "Paso 7 — Top-K: se ordenan los candidatos por score_final descendente y se "
        "devuelven los 10 mejores (TOPK_FINAL = 10)."
    )
    doc.add_paragraph("")

    add_heading(doc, "Ejemplo numérico de score_final", level=2)
    add_paragraph(doc,
        "Supón dos productos candidatos para el cliente u = 'CLI_001':"
    )
    add_table(doc,
        headers=["Score", "Producto A (Queso próximo a vencer)", "Producto B (Bebida nueva)"],
        rows=[
            ["score_cf",       "0.75 (cliente compra mucho queso)",  "0.40 (no es su perfil habitual)"],
            ["score_cbf",      "0.88 (similar a sus compras)",        "0.20 (poco similar)"],
            ["score_urgency",  "0.42 (vence en 10 días)",             "0.01 (vence en 200 días)"],
            ["score_rotation", "0.60 (rotación media-baja)",          "0.95 (baja rotación)"],
            ["score_novelty",  "0.05 (lleva 1 año en catálogo)",      "0.85 (ingresó hace 5 días)"],
            ["score_final",
             "0.40×0.75 + 0.15×0.88 + 0.20×0.42 + 0.15×0.60 + 0.10×0.05\n= 0.300+0.132+0.084+0.090+0.005 = 0.611",
             "0.40×0.40 + 0.15×0.20 + 0.20×0.01 + 0.15×0.95 + 0.10×0.85\n= 0.160+0.030+0.002+0.143+0.085 = 0.420"],
        ],
        col_widths=[1.2, 2.6, 2.6],
    )
    add_paragraph(doc,
        "El Producto A (score_final = 0.611) se recomienda antes que el Producto B (0.420), "
        "porque aunque B es nuevo, el perfil del cliente y la urgencia de A lo hacen más relevante."
    )
    doc.add_paragraph("")

    # ── 9. Conexión con los endpoints ────────────────────────────────────────
    add_heading(doc, "9. Cómo los parámetros matemáticos definen cada endpoint")
    add_paragraph(doc,
        "Los endpoints no cambian la fórmula del score_final. "
        "Lo que cambian es el pool de candidatos sobre el que se aplica:"
    )
    add_table(doc,
        headers=["Endpoint", "Restricción al pool de candidatos", "Efecto matemático"],
        rows=[
            ["/recomendar",          "Ninguna — todos los productos",
             "score_final libre. Gana el mejor equilibrio de los 5 componentes."],
            ["/proximos-vencer",     "Solo productos con días_para_vencer ≤ 30",
             "score_urgency siempre será alto. score_cf y score_cbf discriminan entre urgentes."],
            ["/baja-rotacion",       "Solo productos con baja_rotacion == 1",
             "score_rotation siempre será alto. score_cf discrimina cuál le interesa al cliente."],
            ["/nuevos",              "Solo productos con días_desde_ingreso ≤ 365",
             "score_novelty siempre será alto. score_cbf garantiza coherencia con el perfil."],
            ["/dashboard",           "Los 3 filtros anteriores, sin duplicados",
             "Cada producto aparece una sola vez, en la categoría de mayor prioridad de negocio."],
        ],
        col_widths=[1.5, 2.0, 3.0],
    )
    doc.add_paragraph("")

    output = DOCS_DIR / "modelo_guia_matematica.docx"
    doc.save(output)
    print(f"Generado: {output}")


# ──────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    crear_doc_simple()
    crear_doc_matematico()
    print("Documentación generada correctamente.")
